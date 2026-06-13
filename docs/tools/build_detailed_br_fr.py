from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_srs import (
    ROOT, NAVY, BLUE, LIGHT_BLUE, LIGHT_GRAY, PALE_BLUE, GRAY, GREEN, AMBER,
    RED, PURPLE, BLACK, WHITE, STATUS, requirements, add_page_number,
    add_heading, add_body, add_bullet, add_table, add_status_badge, set_run,
)

OUT = ROOT / "docs" / "DETAILED_BUSINESS_RULES_AND_FUNCTIONAL_REQUIREMENTS.docx"


modules = [
    {
        "code": "AUTH",
        "name": "Authentication & Account Management",
        "theory": [
            "Nguyên tắc định danh duy nhất: một tài khoản phải có định danh ổn định; email và username là thuộc tính đăng nhập, không phải khóa nghiệp vụ thay thế UUID.",
            "Nguyên tắc least privilege: tài khoản mới nhận quyền thấp nhất; quyền cao hơn chỉ do Admin cấp.",
            "Nguyên tắc secure credential lifecycle: mật khẩu phải được băm, token có thời hạn, reset token dùng một lần và không tiết lộ việc email có tồn tại.",
            "Nguyên tắc separation of duties: người quản trị tài khoản không được biết mật khẩu hiện tại của người dùng.",
        ],
        "operations": [
            {
                "id": "AUTH-O1", "title": "Admin tạo tài khoản",
                "status": "IMPLEMENTED",
                "questions": [
                    "Ai được phép tạo tài khoản: Admin, người dùng tự đăng ký, hay cả hai?",
                    "Tài khoản được tạo bằng email thật, email trường hay email giả lập?",
                    "Email và username có bắt buộc duy nhất, không phân biệt hoa thường không?",
                    "Admin tự nhập mật khẩu, hệ thống sinh mật khẩu tạm hay gửi link kích hoạt?",
                    "Tài khoản mới active ngay hay pending chờ xác minh email?",
                    "Admin có được tạo một Admin khác không? Có giới hạn số lượng/quyền không?",
                    "Có hỗ trợ import Excel hàng loạt và gửi email mời không?",
                    "Nếu tạo một phần thành công rồi lỗi thì rollback như thế nào?",
                ],
                "decisions": [
                    "Baseline cho phép tự đăng ký Student và Admin tạo tài khoản cho cả Student/Lecturer/Admin.",
                    "UUID là định danh chính; email và username phải duy nhất không phân biệt hoa thường.",
                    "Mật khẩu không được gửi lại qua API hoặc hiển thị sau khi tạo.",
                    "Tài khoản mới mặc định Student nếu không chỉ định role; role cao hơn phải do Admin xác lập.",
                    "Thiết kế mục tiêu ưu tiên link kích hoạt hoặc mật khẩu tạm bắt buộc đổi; code hiện cho Admin nhập mật khẩu.",
                    "Import Excel và gửi thư mời chưa thuộc baseline, phải ghi là yêu cầu tương lai nếu cần.",
                ],
                "brs": [
                    ("BR-AUTH-01", "Mỗi tài khoản có một UUID bất biến; email/username có thể đổi nhưng không làm đổi ownership dữ liệu.", "IMPLEMENTED"),
                    ("BR-AUTH-02", "Email và username phải duy nhất sau khi trim và chuẩn hóa chữ thường để so sánh.", "PARTIAL"),
                    ("BR-AUTH-03", "Tài khoản tự đăng ký luôn nhận role Student; chỉ Admin được cấp Lecturer hoặc Admin.", "IMPLEMENTED"),
                    ("BR-AUTH-04", "Mật khẩu tối thiểu 8 ký tự trong yêu cầu production; code hiện cho phép tối thiểu 6 ký tự.", "PARTIAL"),
                    ("BR-AUTH-05", "Không actor nào được đọc mật khẩu rõ hoặc password hash của người dùng.", "IMPLEMENTED"),
                    ("BR-AUTH-06", "Tạo tài khoản và credential phải là một transaction; lỗi ở bước nào phải rollback toàn bộ.", "IMPLEMENTED"),
                    ("BR-AUTH-07", "Không được xóa Admin cuối cùng hoặc tự hạ quyền/xóa chính mình nếu làm hệ thống mất quyền quản trị.", "PLANNED"),
                ],
                "fr_ids": ["FR-AUTH-01", "FR-USER-01"],
                "tests": [
                    "Tạo tài khoản hợp lệ với role Student/Lecturer.",
                    "Từ chối email trùng khác hoa thường và username trùng có khoảng trắng.",
                    "Từ chối role không thuộc 1/2/3.",
                    "Mô phỏng lỗi khi insert credential và xác nhận users không có bản ghi mồ côi.",
                    "Response không chứa password/passwordHash.",
                ],
            },
            {
                "id": "AUTH-O2", "title": "Đăng nhập",
                "status": "PARTIAL",
                "questions": [
                    "Cho đăng nhập bằng email, username hay cả hai?",
                    "Tài khoản khóa/inactive có đăng nhập được không?",
                    "Sai mật khẩu bao nhiêu lần thì khóa tạm? Thời gian khóa bao lâu?",
                    "Thông báo lỗi có phân biệt sai email và sai mật khẩu không?",
                    "Token được lưu ở localStorage hay cookie HttpOnly?",
                    "Một người dùng được có bao nhiêu phiên đồng thời?",
                    "Phiên có hết hạn tuyệt đối và hết hạn do không hoạt động không?",
                ],
                "decisions": [
                    "Cho phép email hoặc username, nhưng phản hồi lỗi phải chung để tránh dò tài khoản.",
                    "Tài khoản is_blocked=true hoặc is_active=false phải bị từ chối trước khi phát JWT.",
                    "Production ưu tiên secure HttpOnly SameSite cookie; bearer token chỉ dùng cho Go API với thời hạn ngắn.",
                    "Phải có rate limit theo IP và định danh; sau nhiều lần sai có backoff/lock tạm.",
                    "Role phải lấy từ dữ liệu server và ký trong JWT, không suy đoán từ email.",
                ],
                "brs": [
                    ("BR-AUTH-08", "Thông báo đăng nhập thất bại không tiết lộ email/username có tồn tại.", "PLANNED"),
                    ("BR-AUTH-09", "User blocked hoặc inactive không được tạo session/JWT mới.", "PARTIAL"),
                    ("BR-AUTH-10", "Role trong JWT chỉ được lấy từ DB và chỉ nhận admin/lecturer/student.", "IMPLEMENTED"),
                    ("BR-AUTH-11", "JWT phải có sub, role, iat, exp; production nên có iss, aud và jti.", "PARTIAL"),
                    ("BR-AUTH-12", "Frontend không được suy role từ chuỗi email hoặc tin payload chưa xác minh.", "PARTIAL"),
                    ("BR-AUTH-13", "Sau tối đa 5 lần sai trong 15 phút, hệ thống phải throttle hoặc khóa tạm.", "PLANNED"),
                ],
                "fr_ids": ["FR-AUTH-02", "FR-RBAC-01"],
                "tests": [
                    "Đăng nhập bằng email và username.",
                    "Sai password/email cho cùng một thông báo chung.",
                    "Blocked user không nhận session/JWT.",
                    "JWT hết hạn bị Go trả 401.",
                    "Sửa role client-side không vượt qua RequireRoles.",
                ],
            },
            {
                "id": "AUTH-O3", "title": "Quên và đặt lại mật khẩu",
                "status": "UI",
                "questions": [
                    "Ai được yêu cầu reset: user, Admin hay cả hai?",
                    "Gửi link hay mã OTP? Thời hạn bao lâu?",
                    "Link dùng một lần hay có thể dùng lại?",
                    "Đổi mật khẩu có hủy toàn bộ phiên đang hoạt động không?",
                    "Có lưu lịch sử mật khẩu để cấm dùng lại không?",
                    "Admin reset có được biết mật khẩu mới không?",
                ],
                "decisions": [
                    "User reset qua link ký/mã ngẫu nhiên dùng một lần; phản hồi luôn giống nhau dù email không tồn tại.",
                    "Token hết hạn sau 15-30 phút và bị vô hiệu ngay khi sử dụng hoặc khi phát token mới.",
                    "Đổi mật khẩu thành công phải revoke các phiên khác.",
                    "Admin có thể đặt mật khẩu tạm, nhưng phải buộc người dùng đổi ở lần đăng nhập tiếp theo.",
                ],
                "brs": [
                    ("BR-AUTH-14", "Reset token là bí mật dùng một lần, lưu dạng hash và hết hạn tối đa 30 phút.", "PLANNED"),
                    ("BR-AUTH-15", "Yêu cầu quên mật khẩu luôn trả thông báo trung tính.", "PLANNED"),
                    ("BR-AUTH-16", "Đặt lại mật khẩu thành công phải revoke session/JWT refresh hiện hữu.", "PLANNED"),
                    ("BR-AUTH-17", "Mật khẩu mới không được trùng mật khẩu hiện tại và nên không trùng 5 mật khẩu gần nhất.", "PLANNED"),
                ],
                "fr_ids": ["FR-AUTH-04", "FR-USER-01"],
                "tests": [
                    "Email tồn tại/không tồn tại trả response tương đương.",
                    "Token hết hạn, dùng lại hoặc bị thay thế đều bị từ chối.",
                    "Phiên cũ không dùng được sau reset.",
                    "Admin không nhận lại mật khẩu rõ trong response/log.",
                ],
            },
            {
                "id": "AUTH-O4", "title": "Đăng xuất và kiểm soát phiên",
                "status": "IMPLEMENTED",
                "questions": [
                    "Đăng xuất chỉ phiên hiện tại hay tất cả thiết bị?",
                    "JWT bearer có được revoke trước khi hết hạn không?",
                    "Cookie có Secure, HttpOnly, SameSite phù hợp không?",
                    "Session hết hạn có tự chuyển về login và giữ return URL không?",
                ],
                "decisions": [
                    "Baseline có logout phiên hiện tại; mục tiêu có thêm logout all devices.",
                    "Access token ngắn hạn; refresh/session token được lưu server-side và có thể revoke.",
                    "Mọi client state chứa token phải bị xóa khi logout/401.",
                ],
                "brs": [
                    ("BR-AUTH-18", "Logout phải hủy session server-side và xóa token/cookie phía client.", "IMPLEMENTED"),
                    ("BR-AUTH-19", "Access token production có TTL ngắn; session/refresh token có cơ chế revoke.", "PARTIAL"),
                    ("BR-AUTH-20", "Route guard chỉ hỗ trợ UX; API server luôn là điểm cưỡng chế quyền cuối cùng.", "IMPLEMENTED"),
                ],
                "fr_ids": ["FR-AUTH-03", "FR-RBAC-01"],
                "tests": ["Logout rồi truy cập route bảo vệ.", "Xóa cookie nhưng giữ local token và ngược lại.", "Token role khác truy cập cross-role URL/API."],
            },
        ],
    },
    {
        "code": "DOC",
        "name": "Document Management & Indexing",
        "theory": [
            "Upload là giao dịch nhiều giai đoạn: validate, lưu metadata, lưu object, enqueue, parse, chunk, embed và publish. Mỗi giai đoạn phải có trạng thái và khả năng phục hồi.",
            "Tệp người dùng là dữ liệu không tin cậy: phải kiểm tra kích thước, extension, MIME/magic bytes, tên tệp, nội dung parser và quyền sở hữu.",
            "Tài liệu chỉ có thể phục vụ người đọc/RAG khi đã hoàn tất indexing và thỏa chính sách visibility/approval.",
            "Xóa tài liệu phải có semantic rõ: soft delete để audit/khôi phục hoặc hard delete có cascade và xóa object/vector.",
        ],
        "operations": [
            {
                "id": "DOC-O1", "title": "Upload tài liệu",
                "status": "IMPLEMENTED",
                "questions": [
                    "Tài liệu có bắt buộc liên quan học thuật không? Ai xác nhận?",
                    "Định dạng hỗ trợ chính thức: PDF, DOC/DOCX, PPT/PPTX, TXT, MD?",
                    "Giới hạn dung lượng, số trang và tổng dung lượng mỗi user là bao nhiêu?",
                    "File 0 byte, file đổi extension, file có macro/malware hoặc PDF scan xử lý ra sao?",
                    "Upload nhiều file cùng lúc có được không? Một file lỗi có rollback cả batch không?",
                    "File trùng theo MD5/SHA-256 có cho upload lại không?",
                    "Tệp được lưu ở S3 trước hay tạo record DB trước?",
                    "Document mới pending duyệt hay tự approved sau indexing?",
                ],
                "decisions": [
                    "Lecturer là actor upload; nội dung phải thuộc môn được phân công.",
                    "Baseline cam kết PDF, DOCX, PPTX; DOC/PPT cũ chỉ được nhận khi có converter/parser rõ. TXT/MD phải bổ sung trước khi công bố hỗ trợ.",
                    "Giới hạn mặc định 50 MB mỗi file; file 0 byte bị từ chối.",
                    "Kiểm tra cả extension và MIME/magic bytes; scan PDF không OCR và phải báo không trích xuất được.",
                    "Duplicate hash không tạo document mới trừ khi user xác nhận versioning.",
                    "Document ở processing trong lúc worker chạy; chỉ completed/approved mới xuất hiện cho người khác.",
                ],
                "brs": [
                    ("BR-DOC-01", "Chỉ Lecturer active và được phân công môn mới được upload vào môn đó.", "PARTIAL"),
                    ("BR-DOC-02", "File phải lớn hơn 0 và không vượt MAX_FILE_SIZE=50 MB.", "PARTIAL"),
                    ("BR-DOC-03", "Extension và MIME/magic bytes phải cùng thuộc allowlist.", "PARTIAL"),
                    ("BR-DOC-04", "Tên tệp gốc chỉ dùng hiển thị; storage key phải do server sinh để chống path traversal.", "IMPLEMENTED"),
                    ("BR-DOC-05", "Cùng checksum trong cùng phạm vi owner/môn phải bị chặn hoặc tạo version có chủ ý.", "PARTIAL"),
                    ("BR-DOC-06", "Upload lỗi trước khi enqueue phải bù/xóa metadata và object đã tạo.", "IMPLEMENTED"),
                    ("BR-DOC-07", "PDF scan không có text phải chuyển failed với lý do 'OCR not supported', không tạo chunks rỗng.", "PARTIAL"),
                    ("BR-DOC-08", "Visibility chỉ nhận public, school_wide hoặc private.", "IMPLEMENTED"),
                ],
                "fr_ids": ["FR-DOC-01", "FR-DOC-02", "FR-IDX-04"],
                "tests": [
                    "PDF/DOCX/PPTX hợp lệ; file 0 byte; >50 MB; đổi .exe thành .pdf.",
                    "Lecturer upload môn không được phân công.",
                    "Duplicate checksum cùng owner và khác owner.",
                    "S3 lỗi, DB lỗi, enqueue lỗi và xác nhận không có dữ liệu mồ côi.",
                    "PDF scan chỉ ảnh trả lỗi có nghĩa.",
                ],
            },
            {
                "id": "DOC-O2", "title": "Parse, chunk, embed và phân chương",
                "status": "IMPLEMENTED",
                "questions": [
                    "Chunk size cố định hay theo cấu trúc? Overlap bao nhiêu?",
                    "Nếu một batch embedding lỗi giữa chừng, giữ chunks đã lưu hay rollback?",
                    "Retry bao nhiêu lần và có đổi API key không?",
                    "Embedding dimension nào là chuẩn: 768 hay 3072?",
                    "Phân chương AI lỗi có làm indexing thất bại không?",
                    "Có cho re-index khi đổi model/chunking không? Version cũ xử lý thế nào?",
                    "Worker crash giữa job có được resume hay job bị kẹt processing?",
                ],
                "decisions": [
                    "Chunk mục tiêu 500 token, overlap 100, tôn trọng separator tự nhiên.",
                    "Embedding chuẩn hiện hành là 3072 chiều; mọi tài liệu ghi 768 phải được xem là lỗi thời.",
                    "Phân chương là bước tăng cường: lỗi phải fallback một chương, không hủy chunks đã index.",
                    "Job phải idempotent; re-index phải xóa/thay thế vector theo version, không nhân đôi chunk.",
                    "Job processing quá timeout phải được reclaim hoặc chuyển failed để retry.",
                ],
                "brs": [
                    ("BR-IDX-01", "Chunk không rỗng, có thứ tự liên tục và thuộc đúng một document.", "IMPLEMENTED"),
                    ("BR-IDX-02", "Chunk size mặc định 500 token và overlap 100; thay đổi phải là cấu hình versioned.", "IMPLEMENTED"),
                    ("BR-IDX-03", "Embedding phải đúng 3072 chiều trước khi insert pgvector.", "IMPLEMENTED"),
                    ("BR-IDX-04", "Retry Gemini chỉ áp dụng lỗi tạm thời 429/5xx/network; lỗi 4xx cấu hình phải fail nhanh.", "IMPLEMENTED"),
                    ("BR-IDX-05", "Một job chỉ được một worker claim tại một thời điểm.", "PARTIAL"),
                    ("BR-IDX-06", "Re-index phải idempotent và không tạo duplicate chunks.", "PLANNED"),
                    ("BR-IDX-07", "Segmentation lỗi dùng fallback; embedding/parse lỗi làm toàn job failed.", "IMPLEMENTED"),
                    ("BR-IDX-08", "Document chỉ chuyển completed khi file metadata và toàn bộ chunks đã lưu thành công.", "IMPLEMENTED"),
                ],
                "fr_ids": ["FR-IDX-01", "FR-IDX-02", "FR-IDX-03", "FR-IDX-04", "FR-IDX-05"],
                "tests": [
                    "Parser trả text rỗng/null byte/UTF-8 lỗi.",
                    "Gemini 429, 500, timeout, vector sai chiều.",
                    "Worker crash sau batch đầu rồi chạy lại.",
                    "Segmentation trả JSON lỗi hoặc range chồng lấn.",
                    "Re-index cùng document không tăng gấp đôi chunks.",
                ],
            },
            {
                "id": "DOC-O3", "title": "Xem tài liệu và tài liệu cá nhân",
                "status": "IMPLEMENTED",
                "questions": [
                    "Có xem file gốc hay chỉ nội dung đã extract?",
                    "File processing/failed có hiện trong danh sách nào?",
                    "Sinh viên có thấy tài liệu pending/rejected/private không?",
                    "View count tăng mỗi lần mở, mỗi session hay mỗi user/ngày?",
                    "Tệp đã upload nhưng chưa index có còn trên server không?",
                    "Có hiển thị lỗi indexing và nút retry cho Lecturer không?",
                    "Chapter do AI sinh có được Lecturer sửa không?",
                    "Watermark/download permission áp dụng thế nào?",
                ],
                "decisions": [
                    "Public list chỉ hiển thị tài liệu completed/approved phù hợp visibility.",
                    "Owner/Admin thấy cả processing/failed cùng error và retry action.",
                    "Chi tiết đọc nội dung chunk phân trang; tải/xem file gốc phải dùng presigned URL có thời hạn.",
                    "View count chỉ tăng ở lần mở trang đầu theo cơ chế chống spam tối thiểu.",
                    "Tài liệu chưa index vẫn giữ file để retry trong retention window; quá hạn thì cleanup.",
                ],
                "brs": [
                    ("BR-DOC-09", "Owner/Admin có thể xem trạng thái lỗi; user khác chỉ thấy tài liệu đã publish.", "IMPLEMENTED"),
                    ("BR-DOC-10", "Private chỉ owner/Admin; school_wide cho user xác thực; public theo chính sách triển khai.", "PARTIAL"),
                    ("BR-DOC-11", "Chunk content chỉ trả khi caller có quyền trên document.", "IMPLEMENTED"),
                    ("BR-DOC-12", "Presigned file URL phải ngắn hạn và không lộ S3 credential/key nhạy cảm.", "PARTIAL"),
                    ("BR-DOC-13", "View count không được tăng khi Admin/owner refresh liên tục hoặc crawler gọi.", "PLANNED"),
                    ("BR-DOC-14", "Failed job phải hiển thị lý do an toàn và cho owner retry.", "PARTIAL"),
                ],
                "fr_ids": ["FR-DISC-01", "FR-DISC-02", "FR-DOC-03"],
                "tests": [
                    "Student mở private/pending/rejected của người khác.",
                    "Owner xem failed document và error.",
                    "Chunk pagination biên 1, 8, 10 và vượt tổng trang.",
                    "Presigned URL hết hạn.",
                    "Refresh nhiều lần và kiểm tra view_count.",
                ],
            },
            {
                "id": "DOC-O4", "title": "Sửa, xóa, report và moderation",
                "status": "IMPLEMENTED",
                "questions": [
                    "Sửa metadata có cần duyệt lại không?",
                    "Đổi môn/visibility có ảnh hưởng quyền truy cập và RAG không?",
                    "Xóa là soft delete hay hard delete? Có undo/grace period không?",
                    "Người đang đọc file khi Admin xóa thì sao?",
                    "Xóa cần nhập lại mật khẩu hay chỉ confirm?",
                    "Report trùng cùng user/document có cho phép không?",
                    "Approve/reject có cần lý do và audit log không?",
                ],
                "decisions": [
                    "Chỉ owner sửa; thay đổi metadata nhạy cảm hoặc visibility rộng hơn có thể đưa về pending.",
                    "Baseline hard delete có cascade; production nên soft delete/grace period trước hard cleanup.",
                    "Xóa phải xác nhận rõ tác động; tài liệu đang dùng nhận 404/410 ở request tiếp theo.",
                    "Một user chỉ có một report pending trên cùng document; report bổ sung cập nhật lý do thay vì spam.",
                    "Reject/delete/ignore report phải có resolution note và audit log.",
                ],
                "brs": [
                    ("BR-DOC-15", "Chỉ owner được sửa/xóa tài liệu cá nhân; Admin dùng luồng moderation riêng.", "IMPLEMENTED"),
                    ("BR-DOC-16", "Đổi owner bị cấm; đổi subject phải kiểm tra assignment mới.", "PARTIAL"),
                    ("BR-DOC-17", "Xóa document phải dọn file, chunks, chapters, reports, bookmarks và job theo policy.", "IMPLEMENTED"),
                    ("BR-DOC-18", "Không cho report chính tài liệu của mình trừ khi là yêu cầu gỡ bỏ.", "PLANNED"),
                    ("BR-DOC-19", "Một reporter không tạo nhiều report pending trùng document.", "PLANNED"),
                    ("BR-MOD-01", "Approve/reject/delete/resolve chỉ Admin; mọi thao tác phải audit.", "PARTIAL"),
                    ("BR-MOD-02", "Reject và delete phải có lý do; ignore report phải có resolution note.", "PLANNED"),
                ],
                "fr_ids": ["FR-DOC-03", "FR-MOD-01", "FR-MOD-02"],
                "tests": [
                    "Lecturer sửa/xóa tài liệu người khác.",
                    "Đổi subject sang môn không được phân công.",
                    "Xóa khi S3 lỗi hoặc FK lỗi.",
                    "Report trùng; report own document.",
                    "Admin reject/delete không nhập lý do.",
                ],
            },
        ],
    },
    {
        "code": "CUR",
        "name": "Curriculum, Metadata & Lecturer Assignment",
        "theory": [
            "Master data phải có định danh ổn định, code duy nhất, vòng đời rõ và không được xóa khi đang được tham chiếu nếu không có chiến lược thay thế.",
            "Phân công Lecturer-Subject là quan hệ quyền; thay đổi phải có hiệu lực nhất quán và được audit.",
        ],
        "operations": [
            {
                "id": "CUR-O1", "title": "Quản lý học kỳ, môn học và metadata",
                "status": "IMPLEMENTED",
                "questions": [
                    "Code môn/ngôn ngữ có duy nhất không phân biệt hoa thường không?",
                    "Có được xóa danh mục đang được tài liệu dùng không?",
                    "Học kỳ đóng có cho upload mới không?",
                    "Tên danh mục có hỗ trợ đa ngôn ngữ không?",
                ],
                "decisions": [
                    "Code phải duy nhất sau normalize; name có thể trùng nếu code khác.",
                    "Danh mục đang được tham chiếu không hard delete; dùng inactive hoặc migrate reference.",
                    "Học kỳ có trạng thái planned/active/closed; chỉ active nhận upload mới.",
                ],
                "brs": [
                    ("BR-CUR-01", "Subject code và language code duy nhất không phân biệt hoa thường.", "IMPLEMENTED"),
                    ("BR-CUR-02", "Không hard delete master data đang được tham chiếu.", "PARTIAL"),
                    ("BR-CUR-03", "Chỉ học kỳ active cho phép tạo mới nội dung học thuật.", "PLANNED"),
                    ("BR-CUR-04", "Mọi thay đổi master data phải được audit.", "PARTIAL"),
                ],
                "fr_ids": ["FR-CUR-01", "FR-CUR-02"],
                "tests": ["Code trùng khác hoa thường.", "Xóa subject có document.", "Upload vào term closed.", "Concurrent update cùng metadata."],
            },
            {
                "id": "CUR-O2", "title": "Phân công Lecturer theo môn",
                "status": "IMPLEMENTED",
                "questions": [
                    "Một Lecturer được dạy nhiều môn và một môn có nhiều Lecturer không?",
                    "Gỡ phân công có làm mất quyền với tài liệu đã upload không?",
                    "Admin thay toàn bộ danh sách hay thêm/xóa từng môn?",
                ],
                "decisions": [
                    "Quan hệ nhiều-nhiều; cặp user-subject duy nhất.",
                    "Gỡ assignment ngăn upload mới nhưng không tự đổi ownership tài liệu cũ.",
                    "Replace assignment phải transaction và kiểm tra user role Lecturer.",
                ],
                "brs": [
                    ("BR-CUR-05", "Chỉ user role Lecturer được có user_subjects.", "IMPLEMENTED"),
                    ("BR-CUR-06", "Mỗi cặp Lecturer-Subject là duy nhất.", "IMPLEMENTED"),
                    ("BR-CUR-07", "Gỡ assignment không xóa tài liệu cũ; quyền sửa tài liệu cũ theo ownership và policy.", "PLANNED"),
                    ("BR-CUR-08", "Replace assignment là thao tác nguyên tử.", "PARTIAL"),
                ],
                "fr_ids": ["FR-CUR-03"],
                "tests": ["Assign Student/Admin.", "Cặp trùng.", "Replace bị lỗi giữa chừng.", "Gỡ assignment rồi sửa/upload."],
            },
        ],
    },
    {
        "code": "RAG",
        "name": "Chat Sessions, Retrieval, Generation & Citations",
        "theory": [
            "RAG đáng tin cậy cần tách bốn bước: xác định phạm vi, retrieval, generation có ràng buộc và citation/provenance.",
            "Không tìm thấy bằng chứng là một kết quả hợp lệ; hệ thống phải từ chối thay vì dùng kiến thức nền không kiểm chứng.",
            "Conversation memory phải bị giới hạn theo token, user ownership và không được làm nhiễm truy vấn giữa các môn/tài liệu.",
            "Mọi citation phải truy nguyên được đến chunk thực sự có trong context gửi LLM.",
        ],
        "operations": [
            {
                "id": "RAG-O1", "title": "Tạo và quản lý phiên chat",
                "status": "UI",
                "questions": [
                    "Session gắn với một môn hay có thể chọn nhiều môn/tài liệu?",
                    "Tên session do user nhập hay AI tự sinh?",
                    "Lịch sử lưu bao lâu? Xóa có xóa message/citation không?",
                    "Starred/done có ý nghĩa nghiệp vụ gì?",
                    "Có cho Lecturer xem chat của Student không?",
                ],
                "decisions": [
                    "Session thuộc đúng một user; mặc định một subject, có thể scope thêm document IDs thuộc subject đó.",
                    "Tiêu đề có thể tự sinh từ câu đầu nhưng user được sửa.",
                    "Không role nào xem chat riêng của user khác nếu không có consent/chính sách rõ.",
                    "Xóa session cascade messages/citations theo retention; audit chỉ giữ metadata tối thiểu.",
                ],
                "brs": [
                    ("BR-CHAT-01", "Mỗi session thuộc một user và chỉ owner truy cập.", "NOT"),
                    ("BR-CHAT-02", "Document scope phải là tập con tài liệu user được phép xem.", "NOT"),
                    ("BR-CHAT-03", "Session done là read-only cho đến khi reopen; starred chỉ ảnh hưởng tổ chức danh sách.", "PLANNED"),
                    ("BR-CHAT-04", "Xóa session phải cascade message/citation và không xóa documents/chunks.", "NOT"),
                ],
                "fr_ids": ["FR-CHAT-01"],
                "tests": ["Đọc session người khác.", "Scope private document của người khác.", "Gửi message vào done session.", "Delete cascade."],
            },
            {
                "id": "RAG-O2", "title": "Retrieval ngữ nghĩa",
                "status": "NOT",
                "questions": [
                    "Top-K và similarity threshold bao nhiêu?",
                    "Lọc theo subject/document/chapter trước hay sau vector search?",
                    "Tài liệu pending/private/rejected có được retrieve không?",
                    "Query quá ngắn, rỗng hoặc prompt injection xử lý ra sao?",
                    "Có hybrid search keyword + vector không?",
                ],
                "decisions": [
                    "Filter quyền và trạng thái trước retrieval; không bao giờ dựa vào LLM để lọc quyền.",
                    "Baseline top-K 5-8, cosine threshold cấu hình và có rerank/deduplicate theo document/page.",
                    "Query rỗng bị từ chối; query dài bị giới hạn; prompt injection trong tài liệu/query được coi là nội dung, không phải lệnh hệ thống.",
                    "Nếu không chunk nào đạt threshold, trả out-of-scope, không gọi generation hoặc gọi prompt từ chối có kiểm soát.",
                ],
                "brs": [
                    ("BR-RAG-01", "Chỉ completed/approved và accessible documents tham gia retrieval.", "NOT"),
                    ("BR-RAG-02", "Authorization filter phải áp dụng trước vector similarity.", "NOT"),
                    ("BR-RAG-03", "Không có chunk đạt threshold thì answer.out_of_scope=true.", "NOT"),
                    ("BR-RAG-04", "Top-K, threshold và model là config server-side có version.", "NOT"),
                    ("BR-RAG-05", "Context phải deduplicate chunk overlap và giữ provenance.", "NOT"),
                ],
                "fr_ids": ["FR-RAG-01"],
                "tests": ["Query không context.", "Private/pending chunk có score cao.", "Duplicate overlap.", "Empty/oversized query.", "Cross-subject leakage."],
            },
            {
                "id": "RAG-O3", "title": "Sinh câu trả lời và citation",
                "status": "NOT",
                "questions": [
                    "LLM được dùng kiến thức ngoài context không?",
                    "Nếu nguồn mâu thuẫn nhau thì trả lời thế nào?",
                    "Citation gắn theo câu, đoạn hay toàn answer?",
                    "Regenerate có ghi version không?",
                    "Feedback thumbs up/down được dùng vào đâu?",
                ],
                "decisions": [
                    "Answer phải dựa trên context; kiến thức ngoài chỉ được nêu rõ là không có trong tài liệu nếu policy cho phép.",
                    "Nguồn mâu thuẫn phải nêu mâu thuẫn và trích cả hai, không tự chọn không có căn cứ.",
                    "Citation phải trỏ chunk cụ thể, hiển thị document/chapter/page/excerpt/score.",
                    "Regenerate tạo answer version mới và giữ prompt/context version để audit.",
                    "Feedback lưu user, message, rating, optional reason; không tự thay đổi model trực tiếp.",
                ],
                "brs": [
                    ("BR-RAG-06", "LLM không được xem nội dung tài liệu là system instruction.", "NOT"),
                    ("BR-RAG-07", "Mọi khẳng định dựa nguồn phải có ít nhất một citation hợp lệ.", "NOT"),
                    ("BR-RAG-08", "Citation chỉ được tạo từ chunk đã gửi trong context của answer đó.", "NOT"),
                    ("BR-RAG-09", "Nguồn mâu thuẫn phải được trình bày minh bạch.", "NOT"),
                    ("BR-RAG-10", "Regenerate không ghi đè answer cũ; phải lưu version/parent.", "NOT"),
                    ("BR-RAG-11", "Không lưu/log secret hoặc toàn bộ prompt chứa dữ liệu vượt quyền.", "NOT"),
                ],
                "fr_ids": ["FR-RAG-02", "FR-RAG-03", "FR-RAG-04"],
                "tests": ["Prompt injection trong chunk.", "Citation không nằm trong context.", "Nguồn mâu thuẫn.", "Gemini timeout/empty answer.", "Regenerate và feedback persistence."],
            },
        ],
    },
    {
        "code": "QUIZ",
        "name": "Quiz & Practice",
        "theory": [
            "Đánh giá học tập cần validity, fairness và traceability: câu hỏi phải bám mục tiêu học tập và nguồn; đáp án không lộ trước khi nộp.",
            "Attempt là bản ghi bất biến sau submit; chấm lại phải tạo audit/version nếu đáp án chuẩn thay đổi.",
        ],
        "operations": [
            {
                "id": "QUIZ-O1", "title": "Lecturer sinh và công khai quiz",
                "status": "UI",
                "questions": [
                    "Quiz sinh từ document, chapter hay subject?",
                    "Độ khó được xác định thế nào?",
                    "AI có thể publish thẳng không?",
                    "Mỗi câu cần citation/giải thích không?",
                    "Quiz có lịch mở/đóng, số lần làm và random câu/đáp án không?",
                ],
                "decisions": [
                    "AI chỉ tạo draft; Lecturer phải review và publish.",
                    "Mỗi câu có source chunk, đáp án đúng, giải thích và difficulty.",
                    "Quiz có start/end, duration, max attempts và trạng thái draft/published/closed.",
                    "Sửa quiz đã có attempt phải tạo version mới hoặc hạn chế trường được sửa.",
                ],
                "brs": [
                    ("BR-QUIZ-01", "Chỉ Lecturer được phân công môn mới tạo/publish quiz của môn.", "NOT"),
                    ("BR-QUIZ-02", "AI-generated quiz luôn bắt đầu ở draft.", "NOT"),
                    ("BR-QUIZ-03", "Mỗi câu AI phải có citation đến chunk nguồn.", "NOT"),
                    ("BR-QUIZ-04", "Quiz không publish khi rỗng, thiếu đáp án đúng hoặc ngoài lịch hợp lệ.", "NOT"),
                    ("BR-QUIZ-05", "Quiz đã có attempt không được sửa phá vỡ tính công bằng; dùng version mới.", "NOT"),
                ],
                "fr_ids": ["FR-QUIZ-01"],
                "tests": ["Publish trực tiếp từ AI.", "Question không source/đáp án.", "Lecturer sai môn.", "Sửa đáp án sau khi có attempt."],
            },
            {
                "id": "QUIZ-O2", "title": "Student làm bài và xem kết quả",
                "status": "UI",
                "questions": [
                    "Hết giờ xử lý thế nào? Mất mạng có autosave không?",
                    "Cho quay lại câu trước không?",
                    "Khi nào hiện đáp án và giải thích?",
                    "Có chống làm nhiều tab/attempt đồng thời không?",
                ],
                "decisions": [
                    "Server là nguồn thời gian; hết giờ auto-submit.",
                    "Đáp án autosave; một attempt active cho cùng quiz/user.",
                    "Hiện đáp án theo policy sau submit hoặc sau quiz close.",
                    "Điểm tính server-side từ snapshot version quiz.",
                ],
                "brs": [
                    ("BR-QUIZ-06", "Mỗi attempt gắn user và quiz version, có started/submitted timestamps.", "NOT"),
                    ("BR-QUIZ-07", "Server quyết định deadline và score; client không phải nguồn tin cậy.", "NOT"),
                    ("BR-QUIZ-08", "Không lộ correct answer trước thời điểm policy cho phép.", "NOT"),
                    ("BR-QUIZ-09", "Một user không có nhiều active attempt cho cùng quiz nếu policy không cho.", "NOT"),
                    ("BR-QUIZ-10", "Submitted attempt là bất biến; regrade phải audit.", "NOT"),
                ],
                "fr_ids": ["FR-QUIZ-02", "FR-QUIZ-03"],
                "tests": ["Sửa clock client.", "Hai tab submit.", "Mất mạng gần deadline.", "Đọc correctOptionId qua API trước submit.", "Regrade."],
            },
        ],
    },
    {
        "code": "ADM",
        "name": "Administration, Settings, Audit & Evaluation",
        "theory": [
            "Cấu hình vận hành là dữ liệu nhạy cảm; secret không được đi qua UI dưới dạng rõ và thay đổi phải audit/version.",
            "Audit log phải append-only, đủ actor/action/target/time nhưng không chứa credential.",
            "Đánh giá RAG cần test set đại diện, ground truth, source và quy trình chấm tái lập.",
        ],
        "operations": [
            {
                "id": "ADM-O1", "title": "Cấu hình hệ thống và audit",
                "status": "UI",
                "questions": [
                    "Admin được thay đổi model/API key trực tiếp từ UI không?",
                    "Thay chunk size có tự re-index toàn bộ không?",
                    "Ai xem audit log và retention bao lâu?",
                    "Có rollback cấu hình không?",
                ],
                "decisions": [
                    "Secret được quản lý qua secret store/env; UI chỉ hiển thị masked state và rotate action.",
                    "Đổi model/dimension/chunking tạo config version và yêu cầu migration/re-index có kế hoạch.",
                    "Audit log append-only, Admin đọc, retention tối thiểu 180 ngày.",
                    "Cấu hình có validate, preview impact và rollback.",
                ],
                "brs": [
                    ("BR-ADM-01", "API key/DB credential không được trả rõ về browser.", "PLANNED"),
                    ("BR-ADM-02", "Thay đổi config phải có actor, before/after an toàn và timestamp.", "PARTIAL"),
                    ("BR-ADM-03", "Đổi embedding dimension không có hiệu lực nếu schema/index chưa migrate.", "PLANNED"),
                    ("BR-AUD-01", "Audit log append-only và không chứa password/token/secret.", "PARTIAL"),
                    ("BR-AUD-02", "Các thao tác user, role, moderation, assignment, delete và config phải audit.", "PARTIAL"),
                ],
                "fr_ids": ["FR-SET-02", "FR-AUD-01"],
                "tests": ["Inspect network không thấy secret.", "Concurrent config update.", "Dimension mismatch.", "Audit tampering và sensitive-data scan."],
            },
            {
                "id": "ADM-O2", "title": "Bộ 50 câu hỏi và nghiệm thu RAG",
                "status": "NOT",
                "questions": [
                    "50 câu bao phủ chương nào và tỷ lệ loại câu hỏi ra sao?",
                    "Ground truth do ai duyệt? Có citation chuẩn không?",
                    "Câu ngoài phạm vi chiếm bao nhiêu?",
                    "Đánh giá exact match, semantic, citation hay human rating?",
                ],
                "decisions": [
                    "Dataset tối thiểu 50 câu, phủ nhiều chương và gồm fact, explanation, comparison, synthesis, unanswerable.",
                    "Mỗi câu có expected answer, acceptable variants, source chunks/pages và reviewer.",
                    "Ít nhất 10-20% câu là out-of-scope để kiểm tra từ chối.",
                    "Báo cáo gồm answer correctness, groundedness/citation correctness và out-of-scope accuracy; baseline có thể chấm thủ công, không bắt buộc RAGAS.",
                ],
                "brs": [
                    ("BR-EVAL-01", "Dataset tối thiểu 50 câu và không lấy nguyên từ prompt/demo UI.", "NOT"),
                    ("BR-EVAL-02", "Mỗi ground truth có nguồn kiểm chứng và người duyệt.", "NOT"),
                    ("BR-EVAL-03", "Test set phải có câu ngoài phạm vi và nguồn mâu thuẫn.", "NOT"),
                    ("BR-EVAL-04", "Kết quả đánh giá phải ghi model/config/dataset version.", "NOT"),
                ],
                "fr_ids": ["FR-EVAL-01"],
                "tests": ["Schema dataset.", "Coverage theo chương/type.", "Re-run cùng version.", "Manual review disagreement."],
            },
        ],
    },
]


def req_by_id(req_id):
    return next((r for r in requirements if r["id"] == req_id), None)


def add_rule_table(doc, rows):
    data = [(rid, text, STATUS[status][0]) for rid, text, status in rows]
    table = add_table(doc, ["Business Rule ID", "Quy tắc đã chuẩn hóa", "Hiện trạng"], data, [1800, 5860, 1700], font_size=8.6)
    status_color = {v[0]: v[1] for v in STATUS.values()}
    for row in table.rows[1:]:
        label = row.cells[2].text
        color = status_color.get(label)
        if color:
            from build_srs import shade
            shade(row.cells[2], color)
            for run in row.cells[2].paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string(WHITE)
                run.bold = True


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    header = sec.header.paragraphs[0]
    r = header.add_run("SWD392 Chatbot RAG | Detailed BR & FR Specification")
    set_run(r, size=9, color=GRAY)
    add_page_number(sec.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DETAILED BUSINESS RULES\n& FUNCTIONAL REQUIREMENTS")
    set_run(r, size=26, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SWD392 Chatbot RAG / StudyMate")
    set_run(r, size=17, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phân rã theo operation, câu hỏi nghiệp vụ, lý thuyết, quyết định, BR, FR và test scenario")
    set_run(r, size=11, italic=True, color=GRAY)
    doc.add_paragraph()
    add_table(doc, ["Thuộc tính", "Giá trị"], [
        ("Phiên bản", "1.0"),
        ("Ngày baseline", "13/06/2026"),
        ("Nguồn format", "Authentication FR_AUTH_XX.pdf - phân rã operation và câu hỏi cần chốt"),
        ("Nguồn nội dung", "Toàn bộ codebase + nguyên tắc requirements/security/RAG/data integrity"),
        ("Cách đọc", "Nhãn hiện trạng không thay thế yêu cầu mục tiêu; mục PLANNED/NOT vẫn là specification cần triển khai."),
    ], [2200, 7160], header_fill=LIGHT_GRAY, font_size=9.5)
    doc.add_page_break()

    add_heading(doc, "1. Cách xây dựng BR và FR", 1)
    add_body(doc, "PDF mẫu dùng cách đặt câu hỏi quanh từng operation để phát hiện quyết định còn thiếu. Tài liệu này giữ cách tiếp cận đó, nhưng biến câu trả lời thành Business Rule và Functional Requirement có mã, trạng thái và test scenario.")
    add_heading(doc, "1.1 Phân biệt BR và FR", 2)
    add_table(doc, ["Loại", "Câu hỏi mà nó trả lời", "Ví dụ"], [
        ("Business Rule", "Điều gì luôn đúng hoặc bị cấm trong nghiệp vụ?", "Chỉ Lecturer được phân công môn mới được upload vào môn đó."),
        ("Functional Requirement", "Hệ thống phải làm gì để thực thi rule?", "API upload phải kiểm tra role và user_subjects trước khi nhận file."),
        ("Acceptance/Test", "Làm sao chứng minh yêu cầu đúng?", "Lecturer upload môn không được phân công nhận 403."),
    ], [1800, 3800, 3760])
    add_heading(doc, "1.2 Mẫu chuẩn cho mỗi operation", 2)
    for text in [
        "Câu hỏi nghiệp vụ cần chốt.",
        "Quyết định/assumption theo lý thuyết và phạm vi dự án.",
        "Business Rules có mã BR-<DOMAIN>-XX.",
        "Functional Requirements liên quan có mô tả và acceptance criteria.",
        "Edge cases/test scenarios bắt buộc.",
        "Trạng thái triển khai: đã có, một phần, UI prototype, mục tiêu hoặc chưa có.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "1.3 Quy ước trạng thái", 2)
    status_rows = [(v[0], v[2]) for v in STATUS.values()]
    add_table(doc, ["Nhãn", "Ý nghĩa"], status_rows, [2300, 7060])

    add_heading(doc, "2. Ma trận tổng quan module", 1)
    overview = []
    for module in modules:
        ops = len(module["operations"])
        brs = sum(len(o["brs"]) for o in module["operations"])
        frs = len({x for o in module["operations"] for x in o["fr_ids"]})
        overview.append((module["code"], module["name"], ops, brs, frs))
    add_table(doc, ["Module", "Tên", "Operations", "BR", "FR tham chiếu"], overview, [1100, 4260, 1300, 1100, 1600], font_size=9)

    chapter = 3
    all_brs = []
    all_links = []
    for module in modules:
        add_heading(doc, f"{chapter}. {module['name']} ({module['code']})", 1)
        add_heading(doc, f"{chapter}.1 Nền tảng lý thuyết", 2)
        for theory in module["theory"]:
            add_bullet(doc, theory)
        op_index = 2
        for op in module["operations"]:
            add_heading(doc, f"{chapter}.{op_index} {op['id']} - {op['title']}", 2)
            add_status_badge(doc, op["status"], STATUS[op["status"]][2])
            add_heading(doc, "A. Câu hỏi nghiệp vụ cần chốt", 3)
            for q in op["questions"]:
                add_bullet(doc, q)
            add_heading(doc, "B. Quyết định baseline", 3)
            for d in op["decisions"]:
                add_bullet(doc, d)
            add_heading(doc, "C. Business Rules", 3)
            add_rule_table(doc, op["brs"])
            all_brs.extend((rid, module["code"], op["id"], text, status) for rid, text, status in op["brs"])
            add_heading(doc, "D. Functional Requirements liên quan", 3)
            for fr_id in op["fr_ids"]:
                req = req_by_id(fr_id)
                if req:
                    add_heading(doc, f"{req['id']} - {req['title']}", 3)
                    add_status_badge(doc, req["status"], req["evidence"])
                    add_body(doc, req["shall"])
                    p = doc.add_paragraph()
                    r = p.add_run("Acceptance criteria")
                    set_run(r, bold=True, color=NAVY)
                    for criterion in req["acceptance"]:
                        add_bullet(doc, criterion)
                    if req.get("notes"):
                        add_body(doc, f"Độ lệch hiện trạng: {req['notes']}", bold_prefix="Độ lệch hiện trạng:")
                    for rid, _, _ in op["brs"]:
                        all_links.append((rid, fr_id, op["id"]))
                else:
                    add_body(doc, f"{fr_id}: Chưa tìm thấy trong baseline SRS.")
            add_heading(doc, "E. Edge cases và test scenarios", 3)
            for test in op["tests"]:
                add_bullet(doc, test)
            op_index += 1
        chapter += 1

    add_heading(doc, f"{chapter}. Danh mục Business Rules", 1)
    add_table(doc, ["BR ID", "Module", "Operation", "Rule", "Trạng thái"],
              [(a, b, c, d, STATUS[e][0]) for a, b, c, d, e in all_brs],
              [1450, 900, 1300, 4110, 1600], font_size=7.8)
    chapter += 1

    add_heading(doc, f"{chapter}. Danh mục Functional Requirements đầy đủ", 1)
    fr_rows = []
    for req in requirements:
        fr_rows.append((req["id"], req["title"], req["shall"], STATUS[req["status"]][0]))
    add_table(doc, ["FR ID", "Tên", "Yêu cầu hệ thống", "Trạng thái"], fr_rows,
              [1450, 2100, 4210, 1600], font_size=7.8)
    chapter += 1

    add_heading(doc, f"{chapter}. Ma trận truy vết BR - FR - Operation", 1)
    unique_links = []
    seen = set()
    for link in all_links:
        if link not in seen:
            seen.add(link)
            unique_links.append(link)
    add_table(doc, ["Business Rule", "Functional Requirement", "Operation"],
              unique_links, [3000, 3360, 3000], font_size=8.5)
    chapter += 1

    add_heading(doc, f"{chapter}. Các quyết định cần Product Owner xác nhận", 1)
    decisions = [
        ("D-01", "Tài khoản do user tự đăng ký, Admin tạo, hay cả hai?", "Đề xuất: cả hai; self-register chỉ Student."),
        ("D-02", "Mật khẩu tối thiểu 6 hay 8 ký tự?", "Đề xuất production: >=8 và kiểm tra độ mạnh."),
        ("D-03", "DOC/PPT cũ, TXT và MD có thuộc file support chính thức?", "Đề xuất chỉ công bố loại đã có parser/test."),
        ("D-04", "Tài liệu sau indexing tự approved hay chờ Admin duyệt?", "Đề xuất school-wide/public phải duyệt; private có thể tự hoàn tất."),
        ("D-05", "Xóa tài liệu hard delete ngay hay có grace period?", "Đề xuất soft delete 7-30 ngày rồi cleanup."),
        ("D-06", "Chat scope một subject hay đa subject?", "Đề xuất một subject/session, chọn nhiều documents trong subject."),
        ("D-07", "Top-K và similarity threshold mặc định?", "Đề xuất bắt đầu K=6, threshold hiệu chỉnh bằng test set."),
        ("D-08", "Quiz có nằm trong phạm vi release không?", "Hiện chỉ UI prototype; cần quyết định để bổ sung schema/API."),
        ("D-09", "Retention chat, audit, upload jobs và file failed?", "Cần policy chính thức trước production."),
        ("D-10", "Bộ 50 câu do ai sở hữu và duyệt?", "Đề xuất Lecturer/SME soạn, ít nhất một reviewer độc lập."),
    ]
    add_table(doc, ["ID", "Câu hỏi quyết định", "Khuyến nghị"], decisions, [1100, 4300, 3960], font_size=8.7)

    add_heading(doc, "Kết luận", 1)
    add_body(doc, "Codebase hiện có nền tảng tương đối đầy đủ cho auth, RBAC, document management và indexing. Tuy nhiên, Chat/RAG end-to-end, citation persistence, quiz/practice, reset password hoàn chỉnh, bookmark và system settings vẫn thiếu hoặc chỉ là UI prototype. Các Business Rule trong tài liệu này bao gồm cả rule đã hiện thực hóa và rule lý thuyết cần thiết để sản phẩm đúng, an toàn và có thể nghiệm thu.")

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.08

    doc.core_properties.title = "Detailed Business Rules and Functional Requirements - SWD392 Chatbot RAG"
    doc.core_properties.subject = "Operation-based BR and FR specification"
    doc.core_properties.author = "SWD392 Team"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
