from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "SRS_SWD392_Chatbot_RAG.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
GREEN = "2E7D32"
AMBER = "9A6700"
RED = "B42318"
PURPLE = "6941C6"
BLACK = "1F2937"
WHITE = "FFFFFF"

STATUS = {
    "IMPLEMENTED": ("ĐÃ TRIỂN KHAI", GREEN, "Có bằng chứng backend/API hoặc luồng hoạt động thực trong code."),
    "PARTIAL": ("MỘT PHẦN", AMBER, "Có một phần luồng thực, nhưng còn thiếu hành vi, tích hợp hoặc kiểm thử."),
    "UI": ("UI PROTOTYPE", PURPLE, "Có màn hình/interaction nhưng chủ yếu dùng mock, state cục bộ hoặc mô phỏng."),
    "PLANNED": ("YÊU CẦU MỤC TIÊU", BLUE, "Được yêu cầu trong phạm vi sản phẩm nhưng chưa có bằng chứng triển khai đầy đủ."),
    "NOT": ("CHƯA TRIỂN KHAI", RED, "Không tìm thấy API, domain model hoặc luồng thực thi tương ứng."),
    "OOS": ("NGOÀI PHẠM VI", GRAY, "Không thuộc phạm vi phát hành hiện tại."),
}


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=None, bold=None, color=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_status_badge(doc, status_key, evidence):
    label, color, meaning = STATUS[status_key]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2100, 7260])
    set_repeat_table_header(table.rows[0])
    left, right = table.rows[0].cells
    shade(left, color)
    shade(right, PALE_BLUE)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    set_run(r, size=9, bold=True, color=WHITE)
    p2 = right.paragraphs[0]
    r2 = p2.add_run(evidence or meaning)
    set_run(r2, size=9.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_requirement(doc, req):
    add_heading(doc, f"{req['id']} - {req['title']}", 3)
    add_status_badge(doc, req["status"], req["evidence"])
    add_body(doc, req["shall"])
    p = doc.add_paragraph()
    r = p.add_run("Tiêu chí nghiệm thu")
    set_run(r, bold=True, color=NAVY)
    for item in req["acceptance"]:
        add_bullet(doc, item)
    if req.get("notes"):
        p = doc.add_paragraph()
        r = p.add_run("Ghi chú/độ lệch hiện trạng: ")
        set_run(r, bold=True, color=RED)
        r = p.add_run(req["notes"])
        set_run(r, color=BLACK)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, header_fill)
        p = cell.paragraphs[0]
        r = p.add_run(value)
        set_run(r, size=font_size, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_run(r, size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


requirements = [
    {
        "id": "FR-AUTH-01", "title": "Đăng ký tài khoản",
        "status": "IMPLEMENTED",
        "shall": "Hệ thống phải cho phép người dùng đăng ký bằng email, mật khẩu, tên hiển thị và username duy nhất thông qua Better Auth.",
        "evidence": "Better Auth/Hono có endpoint sign-up email, username plugin và bảng users/account.",
        "acceptance": ["Email và username trùng bị từ chối.", "Mật khẩu tối thiểu 6 ký tự.", "Mật khẩu chỉ được lưu dưới dạng băm trong bảng account.", "Tài khoản mới mặc định có vai trò Student (role_id=3)."],
        "notes": "Chưa yêu cầu xác minh email; Better Auth cấu hình requireEmailVerification=false.",
    },
    {
        "id": "FR-AUTH-02", "title": "Đăng nhập và phát hành JWT",
        "status": "PARTIAL",
        "shall": "Hệ thống phải xác thực bằng email hoặc username và mật khẩu, sau đó phát hành JWT có sub, role và thời hạn hết hạn.",
        "evidence": "Better Auth hỗ trợ username plugin và JWT HS256; UI đăng nhập hiện gọi signIn.email.",
        "acceptance": ["Thông tin sai trả về 401/4xx.", "JWT có sub là UUID người dùng và role hợp lệ.", "JWT hết hạn sau thời gian cấu hình.", "Frontend chuyển hướng đúng trang chủ theo vai trò."],
        "notes": "Màn hình hiện nhập email; luồng username login chưa được chứng minh end-to-end. Frontend còn fallback suy vai trò từ email.",
    },
    {
        "id": "FR-AUTH-03", "title": "Đăng xuất và quản lý phiên",
        "status": "IMPLEMENTED",
        "shall": "Hệ thống phải cho phép người dùng đăng xuất, hủy phiên Better Auth và xóa token phía trình duyệt.",
        "evidence": "useAuth.signOut gọi authClient.signOut, xóa localStorage/cookie và chuyển về /login.",
        "acceptance": ["Sau đăng xuất, route được bảo vệ chuyển về /login.", "Token cục bộ bị xóa.", "Phiên phía auth service không còn hợp lệ."],
    },
    {
        "id": "FR-AUTH-04", "title": "Khôi phục mật khẩu",
        "status": "UI",
        "shall": "Hệ thống phải cho phép người dùng yêu cầu liên kết đặt lại mật khẩu qua email và hoàn tất đặt lại bằng token một lần.",
        "evidence": "Có trang forgot-password và lời gọi forgetPassword, nhưng UI hiện có mô phỏng setTimeout và không thấy cấu hình email/reset hoàn chỉnh.",
        "acceptance": ["Không tiết lộ email có tồn tại hay không.", "Token có thời hạn và chỉ dùng một lần.", "Mật khẩu mới đáp ứng chính sách.", "Người dùng có thể đăng nhập bằng mật khẩu mới."],
    },
    {
        "id": "FR-RBAC-01", "title": "Phân quyền ba vai trò",
        "status": "IMPLEMENTED",
        "shall": "Hệ thống phải áp dụng ba vai trò Admin, Lecturer và Student trên cả frontend route guard và API.",
        "evidence": "JWT role, RequireRoles trong Gin router và middleware Next.js theo /admin, /lecturer, /student.",
        "acceptance": ["Không có token trả 401.", "Sai vai trò trả 403 hoặc bị chuyển hướng.", "Student không thể gọi API upload/admin.", "Lecturer không thể gọi API quản trị hệ thống."],
    },
    {
        "id": "FR-USER-01", "title": "Quản lý người dùng",
        "status": "IMPLEMENTED",
        "shall": "Admin phải có thể xem, tạo, cập nhật, khóa/mở khóa, đổi mật khẩu và xóa người dùng.",
        "evidence": "Go cung cấp list/block/unblock; Hono cung cấp create/update/password/delete; frontend có user table và modals.",
        "acceptance": ["Chỉ Admin thực hiện được.", "Không tạo email/username trùng.", "Khóa tài khoản cập nhật is_active/is_blocked.", "Đổi mật khẩu cập nhật bản ghi credential.", "Thao tác lỗi hiển thị thông báo có ý nghĩa."],
        "notes": "Cần kiểm thử hành vi đăng nhập của tài khoản bị khóa; Better Auth config chưa thể hiện hook chặn is_blocked.",
    },
    {
        "id": "FR-CUR-01", "title": "Quản lý học kỳ và môn học",
        "status": "IMPLEMENTED",
        "shall": "Admin phải có thể tạo, sửa, xóa học kỳ và môn học, đồng thời tra cứu danh mục để dùng trong bộ lọc và upload.",
        "evidence": "CRUD academic-terms/subjects trong Go router, repositories và frontend curriculum API.",
        "acceptance": ["Mã môn học duy nhất.", "Môn học có thể liên kết học kỳ.", "Danh mục cập nhật xuất hiện trong lookups.", "Không xóa dữ liệu đang được tham chiếu nếu gây mất toàn vẹn."],
    },
    {
        "id": "FR-CUR-02", "title": "Quản lý loại, nguồn và ngôn ngữ tài liệu",
        "status": "IMPLEMENTED",
        "shall": "Admin phải có thể CRUD loại tài liệu, nguồn tài liệu và ngôn ngữ.",
        "evidence": "Có đầy đủ route, handler, repository và bảng document_types, document_sources, languages.",
        "acceptance": ["Giá trị code/name hợp lệ và không trùng theo ràng buộc DB.", "Lookups phản ánh thay đổi.", "Xóa dữ liệu đang dùng phải bị từ chối hoặc xử lý có kiểm soát."],
    },
    {
        "id": "FR-CUR-03", "title": "Phân công giảng viên theo môn học",
        "status": "IMPLEMENTED",
        "shall": "Admin phải có thể xem và thay thế danh sách môn học được phân công cho từng Lecturer.",
        "evidence": "user_subjects, lecturer-subject repository, GET user-subjects và PUT lecturers/:id/subjects.",
        "acceptance": ["Chỉ người dùng vai trò Lecturer được phân công.", "Không tạo cặp trùng.", "Thao tác thay thế có tính nhất quán.", "Danh sách upload/lọc của Lecturer tôn trọng phân công."],
        "notes": "Cần test rõ việc chặn upload vào môn chưa được phân công.",
    },
    {
        "id": "FR-DOC-01", "title": "Upload tài liệu",
        "status": "IMPLEMENTED",
        "shall": "Lecturer phải có thể upload PDF, DOC, DOCX, PPT hoặc PPTX kèm metadata, visibility và tệp tối đa 50 MB.",
        "evidence": "POST /api/documents/upload, validation extension, S3 storage, upload job và UI upload.",
        "acceptance": ["Thiếu tệp hoặc sai định dạng bị từ chối.", "Tệp vượt MAX_FILE_SIZE bị từ chối.", "Metadata được lưu cùng owner.", "Upload thành công tạo document và job nền.", "Lỗi S3 không để lại document mồ côi."],
        "notes": "Tài liệu định hướng cũ nói hỗ trợ TXT/MD, nhưng handler hiện chỉ chấp nhận PDF/DOC/DOCX/PPT/PPTX.",
    },
    {
        "id": "FR-DOC-02", "title": "Lưu trữ tệp và kiểm tra trùng",
        "status": "PARTIAL",
        "shall": "Hệ thống phải lưu tệp gốc an toàn, ghi checksum và ngăn upload trùng theo chính sách dự án.",
        "evidence": "Có S3FileStorage, checksum SHA-256, md5_hash và luồng mock S3 local.",
        "acceptance": ["Tên tệp không được dùng trực tiếp để tạo đường dẫn nguy hiểm.", "Checksum được ghi lại.", "Tệp trùng được phát hiện theo MD5/checksum.", "Xóa document xóa cả object storage liên quan."],
        "notes": "Chính sách trùng và hành vi local mock cần được kiểm thử tích hợp.",
    },
    {
        "id": "FR-IDX-01", "title": "Trích xuất văn bản",
        "status": "IMPLEMENTED",
        "shall": "Worker phải trích xuất văn bản từ PDF, DOCX, PPTX và văn bản thuần theo parser tương ứng, giữ thông tin trang/slide khi có.",
        "evidence": "Parser factory và parser PDF/DOCX/PPTX/TEXT cùng unit tests.",
        "acceptance": ["Tài liệu rỗng chuyển trạng thái lỗi.", "UTF-8/null bytes được làm sạch.", "Trang hoặc slide được gắn vào kết quả khi parser hỗ trợ.", "Không OCR tài liệu ảnh scan."],
        "notes": "DOC/PPT kiểu cũ được upload cho phép nhưng parser factory cần được xác nhận có xử lý hoặc chuyển đổi.",
    },
    {
        "id": "FR-IDX-02", "title": "Chunking tài liệu",
        "status": "IMPLEMENTED",
        "shall": "Hệ thống phải chia văn bản thành chunk khoảng 500 token với overlap 100 token, ưu tiên ranh giới ngữ nghĩa tự nhiên.",
        "evidence": "TextChunker(500,100), recursive splitter, hashing và bộ unit tests.",
        "acceptance": ["Không tạo chunk rỗng.", "Thứ tự chunk liên tục.", "Overlap không vượt kích thước chunk.", "Nội dung lớn được chia ổn định.", "Hash hỗ trợ nhận diện nội dung trùng."],
    },
    {
        "id": "FR-IDX-03", "title": "Sinh embedding và lưu pgvector",
        "status": "IMPLEMENTED",
        "shall": "Worker phải sinh vector 3072 chiều bằng Gemini Embedding 2 theo batch và lưu vào pgvector.",
        "evidence": "Embedding client yêu cầu outputDimensionality=3072; DB dùng vector(3072); batch 50 và retry/backoff.",
        "acceptance": ["Mỗi chunk thành công có vector 3072 chiều.", "Batch lỗi làm job thất bại có thông báo.", "Rate limit được retry có backoff.", "Không log API key hoặc vector thô."],
        "notes": "Tài liệu ERD/Architecture cũ ghi 768 chiều; code và schema hiện hành dùng 3072.",
    },
    {
        "id": "FR-IDX-04", "title": "Theo dõi tiến độ indexing",
        "status": "IMPLEMENTED",
        "shall": "Hệ thống phải xử lý upload bất đồng bộ và cung cấp trạng thái, phần trăm, thông báo tiến độ cho Lecturer.",
        "evidence": "upload_jobs, worker polling 5 giây, dashboard/my-documents trả active jobs.",
        "acceptance": ["Trạng thái đi qua pending/processing/done hoặc failed.", "Progress từ 0 đến 100.", "Lỗi cập nhật document status=error.", "Job hoàn tất cập nhật tổng chunk/chapter."],
    },
    {
        "id": "FR-IDX-05", "title": "Phân chương bằng AI",
        "status": "IMPLEMENTED",
        "shall": "Sau indexing, hệ thống phải tạo cấu trúc chương và tóm tắt tiếng Việt bằng Gemini, đồng thời có fallback khi API thất bại.",
        "evidence": "GeminiChapterSegmentationService, document_chapters và fallback một chương.",
        "acceptance": ["Các khoảng chunk không chồng lấn ngoài ý muốn.", "Mỗi chương có title/order/range.", "Fallback vẫn cho phép đọc tài liệu.", "Lỗi segmentation không làm mất chunks đã index."],
    },
    {
        "id": "FR-DISC-01", "title": "Danh sách và tìm kiếm tài liệu",
        "status": "IMPLEMENTED",
        "shall": "Student và Lecturer phải xem danh sách tài liệu được phép, tìm kiếm, lọc, sắp xếp và phân trang.",
        "evidence": "GET /documents có q, subject/type/language/source, sortBy, page/pageSize; UI shared documents gọi API nhưng có fallback mock.",
        "acceptance": ["Kết quả tôn trọng visibility/status.", "Phân trang trả tổng số và tổng trang.", "Bộ lọc có thể kết hợp.", "Sắp xếp hỗ trợ title/date/views.", "Không lộ private document của người khác."],
        "notes": "Một số view frontend vẫn chèn mock khi API rỗng hoặc lỗi.",
    },
    {
        "id": "FR-DISC-02", "title": "Đọc chi tiết tài liệu",
        "status": "IMPLEMENTED",
        "shall": "Người dùng được phép phải xem metadata, file, chương và nội dung chunk theo trang bằng slug.",
        "evidence": "GET /documents/:slug với chunkPage/chunkPageSize; details DTO và UI document detail.",
        "acceptance": ["Slug không tồn tại trả 404.", "Không đủ quyền trả 403.", "Chunk được sắp theo thứ tự.", "Lượt xem chỉ tăng theo quy tắc.", "Nội dung dài được phân trang 8-10 chunk/trang."],
    },
    {
        "id": "FR-DISC-03", "title": "Bookmark tài liệu",
        "status": "NOT",
        "shall": "Student và Lecturer phải có thể đánh dấu/bỏ đánh dấu tài liệu và xem danh sách bookmark cá nhân.",
        "evidence": "Schema có user_bookmarks nhưng không tìm thấy repository, API hoặc UI tích hợp hoàn chỉnh.",
        "acceptance": ["Bookmark là duy nhất theo user-document.", "Chỉ tài liệu người dùng được xem mới bookmark được.", "Xóa tài liệu tự dọn bookmark.", "Danh sách bookmark có phân trang."],
    },
    {
        "id": "FR-DOC-03", "title": "Quản lý tài liệu của Lecturer",
        "status": "IMPLEMENTED",
        "shall": "Lecturer phải xem dashboard, danh sách tài liệu cá nhân, sửa metadata, xem tác động xóa và xóa tài liệu thuộc sở hữu.",
        "evidence": "dashboard, my, edit, delete-view, delete routes và service/repository tương ứng.",
        "acceptance": ["Chỉ owner được sửa/xóa.", "Xóa dọn file, chapter, chunk và job liên quan.", "Edit kiểm tra UUID/visibility.", "Dashboard tổng hợp document/file/chunk/job."],
    },
    {
        "id": "FR-MOD-01", "title": "Báo cáo tài liệu",
        "status": "IMPLEMENTED",
        "shall": "Student và Lecturer phải có thể gửi báo cáo vi phạm đối với tài liệu đang xem.",
        "evidence": "POST /documents/:slug/report, document_reports và moderation API.",
        "acceptance": ["Reason bắt buộc.", "Reporter được lấy từ JWT.", "Báo cáo lưu trạng thái pending.", "Không cho báo cáo tài liệu không tồn tại/không được xem."],
    },
    {
        "id": "FR-MOD-02", "title": "Duyệt và xử lý tài liệu",
        "status": "IMPLEMENTED",
        "shall": "Admin phải xem toàn bộ tài liệu, approve, reject hoặc xóa và xử lý báo cáo bằng ignore/delete.",
        "evidence": "Admin document/report routes, service và UI moderation/documents.",
        "acceptance": ["Chỉ Admin thực hiện được.", "Approve cập nhật approved_at.", "Reject cập nhật status.", "Delete dọn tài nguyên.", "Resolve report ghi nhận kết quả."],
        "notes": "Frontend dùng mock fallback khi API lỗi; audit log cho thao tác quản trị cần xác nhận đầy đủ.",
    },
    {
        "id": "FR-CHAT-01", "title": "Tạo và quản lý phiên chat",
        "status": "UI",
        "shall": "Người dùng phải tạo, đổi tên, đánh dấu sao, kết thúc, xóa và xem lịch sử phiên chat của chính mình.",
        "evidence": "UI sessions và sessions-store cục bộ; không có bảng chat_sessions/messages trong schema hiện hành và không có Go chat routes.",
        "acceptance": ["Phiên gắn với user và môn học.", "Người dùng chỉ thấy phiên của mình.", "Lịch sử bền vững sau reload/đăng nhập lại.", "Xóa phiên yêu cầu xác nhận.", "Danh sách hỗ trợ tìm kiếm và lọc starred/status."],
    },
    {
        "id": "FR-RAG-01", "title": "Truy xuất ngữ nghĩa",
        "status": "NOT",
        "shall": "Khi nhận câu hỏi, hệ thống phải embed truy vấn và tìm top-K chunk liên quan bằng cosine similarity trong phạm vi môn/tài liệu người dùng chọn.",
        "evidence": "Có embedding và vector storage cho indexing, nhưng chunk repository chưa có semantic search và router không có query/chat endpoint.",
        "acceptance": ["Chỉ tìm trong tài liệu completed/approved và người dùng được phép xem.", "Hỗ trợ lọc subject/document IDs.", "Trả score và metadata nguồn.", "Ngưỡng thấp dẫn đến out-of-scope.", "Truy vấn được bảo vệ bởi JWT."],
    },
    {
        "id": "FR-RAG-02", "title": "Sinh câu trả lời có căn cứ",
        "status": "NOT",
        "shall": "Hệ thống phải gửi câu hỏi, lịch sử hội thoại giới hạn và context truy xuất đến Gemini LLM để sinh câu trả lời chỉ dựa trên nguồn.",
        "evidence": "Có system prompt và Gemini dùng cho phân chương, nhưng không có chat use case/LLM client/handler thực tế.",
        "acceptance": ["Prompt chống bịa đặt và prompt injection cơ bản.", "Không đủ context phải từ chối có kiểm soát.", "Câu trả lời giữ ngôn ngữ người dùng khi phù hợp.", "Không đưa API key/secret vào prompt.", "Lỗi Gemini trả thông báo có thể thử lại."],
    },
    {
        "id": "FR-RAG-03", "title": "Trích dẫn nguồn",
        "status": "UI",
        "shall": "Mỗi câu trả lời có căn cứ phải kèm trích dẫn tới tài liệu, chương/trang/slide, excerpt và relevance score.",
        "evidence": "UI hiển thị citation card mock; chưa có message_citations schema/API trong DB hiện hành.",
        "acceptance": ["Citation ánh xạ đến chunk thực.", "Nhấp citation mở đoạn nguồn chính xác.", "Không trích dẫn tài liệu người dùng không được xem.", "Citation vẫn tồn tại cùng lịch sử chat."],
    },
    {
        "id": "FR-RAG-04", "title": "Phản hồi streaming và điều khiển chat",
        "status": "UI",
        "shall": "Frontend phải hiển thị trạng thái đang sinh, tự cuộn có kiểm soát, cho phép copy, regenerate và feedback.",
        "evidence": "Có typing simulation, copy/regenerate/feedback controls trong UI; không có streaming transport hoặc API feedback.",
        "acceptance": ["Loading rõ ràng.", "Không gửi lặp khi đang xử lý.", "Người dùng có thể dừng auto-scroll.", "Copy sao chép đúng nội dung.", "Regenerate tạo phiên bản mới có truy vết.", "Feedback được lưu."],
    },
    {
        "id": "FR-QUIZ-01", "title": "Sinh quiz từ tài liệu",
        "status": "UI",
        "shall": "Lecturer phải chọn môn/tài liệu, số câu, độ khó và yêu cầu AI tạo quiz có đáp án để duyệt trước khi công khai.",
        "evidence": "Có create quiz/practice UI và mock generator; không có quiz schema/API/backend.",
        "acceptance": ["Câu hỏi chỉ dựa trên tài liệu được chọn.", "Mỗi câu có đáp án đúng và giải thích/citation.", "Lecturer sửa/xóa câu trước publish.", "Quiz có thời lượng, phạm vi và trạng thái.", "Không publish quiz rỗng."],
    },
    {
        "id": "FR-QUIZ-02", "title": "Làm bài và chấm quiz",
        "status": "UI",
        "shall": "Student phải xem quiz được giao, làm bài trong thời gian quy định, nộp bài và xem kết quả theo chính sách.",
        "evidence": "useQuiz chấm điểm trên state mock; không có persistence hoặc API.",
        "acceptance": ["Đáp án được lưu an toàn.", "Hết giờ tự nộp.", "Mỗi attempt được ghi nhận.", "Điểm được tính nhất quán.", "Không lộ đáp án trước khi nộp."],
    },
    {
        "id": "FR-QUIZ-03", "title": "Lịch sử và phân tích kết quả",
        "status": "UI",
        "shall": "Student phải xem lịch sử điểm; Lecturer phải xem thống kê lớp và chi tiết kết quả từng sinh viên.",
        "evidence": "Có biểu đồ và bảng kết quả mock ở student/lecturer practice.",
        "acceptance": ["Dữ liệu lấy từ attempts thực.", "Phân quyền đúng lớp/môn.", "Hiển thị score, thời gian, đúng/sai, ngày nộp.", "Có thống kê xu hướng và tỷ lệ hoàn thành."],
    },
    {
        "id": "FR-SET-01", "title": "Cài đặt cá nhân",
        "status": "UI",
        "shall": "Người dùng phải cập nhật tên hiển thị, tùy chọn giao diện/ngôn ngữ và các tùy chọn thông báo cá nhân.",
        "evidence": "Có SettingsView nhưng chưa tìm thấy API/profile persistence tương ứng.",
        "acceptance": ["Dữ liệu được lưu theo user.", "Email/username thay đổi phải kiểm tra duy nhất.", "Thay đổi nhạy cảm yêu cầu xác thực lại.", "Không hiển thị secret."],
    },
    {
        "id": "FR-SET-02", "title": "Cấu hình hệ thống",
        "status": "UI",
        "shall": "Admin phải quản lý cấu hình embedding, chunking, ngưỡng retrieval, thông báo và chính sách bảo mật thông qua backend an toàn.",
        "evidence": "Admin settings hiện chỉ là state cục bộ và setTimeout.",
        "acceptance": ["Không trả API key rõ về frontend.", "Thay đổi được audit.", "Giá trị được validate.", "Cấu hình có version/rollback.", "Chỉ Admin truy cập."],
    },
    {
        "id": "FR-AUD-01", "title": "Nhật ký kiểm toán",
        "status": "PARTIAL",
        "shall": "Hệ thống phải ghi audit log cho thao tác quản trị và thay đổi dữ liệu nhạy cảm.",
        "evidence": "Có bảng/domain/repository audit_logs và service dependency, nhưng mức phủ thao tác chưa được chứng minh toàn diện.",
        "acceptance": ["Log có actor, action, target, time, IP khi có.", "Không ghi mật khẩu/token/secret.", "Log không bị người dùng thường sửa.", "Admin có thể tra cứu theo thời gian/actor/action."],
    },
    {
        "id": "FR-EVAL-01", "title": "Bộ kiểm thử 50 câu hỏi",
        "status": "NOT",
        "shall": "Dự án phải cung cấp test set tối thiểu 50 câu hỏi kèm ground truth để đánh giá độ chính xác chatbot.",
        "evidence": "Được nêu trong Requirement.txt nhưng không tìm thấy dataset trong repository.",
        "acceptance": ["Có ít nhất 50 bản ghi.", "Mỗi bản ghi có question, expected answer, source/citation.", "Bao phủ nhiều chương và câu ngoài phạm vi.", "Có quy trình chạy và báo cáo kết quả tái lập."],
        "notes": "Benchmark/RAGAS tự động vẫn ngoài phạm vi; test set thủ công là deliverable bắt buộc.",
    },
]


use_cases = [
    ("UC-01", "Đăng nhập", "Visitor", "Có tài khoản hợp lệ", "Nhập email/username và mật khẩu; hệ thống xác thực, phát JWT, xác định role và chuyển hướng.", "Sai thông tin, tài khoản khóa, token không phát hành.", "IMPLEMENTED"),
    ("UC-02", "Upload và indexing", "Lecturer", "Đã đăng nhập, được phân công môn", "Chọn tệp/metadata; hệ thống lưu S3, tạo job, parse, chunk, embed, phân chương và báo hoàn tất.", "Sai định dạng, quá dung lượng, S3/Gemini/DB lỗi.", "IMPLEMENTED"),
    ("UC-03", "Tìm và đọc tài liệu", "Student/Lecturer", "Đã đăng nhập", "Tìm kiếm/lọc; mở slug; xem metadata, chương và chunk phân trang.", "Không có quyền, slug không tồn tại.", "IMPLEMENTED"),
    ("UC-04", "Quản lý tài liệu cá nhân", "Lecturer", "Sở hữu tài liệu", "Xem dashboard, sửa metadata, xem tác động và xác nhận xóa.", "Tài liệu không thuộc sở hữu, storage delete lỗi.", "IMPLEMENTED"),
    ("UC-05", "Duyệt nội dung", "Admin", "Đã đăng nhập Admin", "Xem tài liệu pending/reports; approve, reject, ignore hoặc delete.", "Đối tượng đã bị xóa, DB/storage lỗi.", "IMPLEMENTED"),
    ("UC-06", "Quản lý người dùng", "Admin", "Đã đăng nhập Admin", "Tạo/sửa/khóa/đổi mật khẩu/xóa user.", "Trùng email/username, xóa vi phạm FK.", "IMPLEMENTED"),
    ("UC-07", "Chat RAG", "Student/Lecturer", "Có tài liệu indexed và phiên chat", "Gửi câu hỏi; hệ thống retrieve top-K, sinh câu trả lời, lưu message và citation.", "Không đủ context, Gemini lỗi, tài liệu bị thu hồi.", "NOT"),
    ("UC-08", "Sinh và làm quiz", "Lecturer/Student", "Có tài liệu approved", "Lecturer sinh/publish quiz; Student làm/nộp; hệ thống chấm và lưu lịch sử.", "AI tạo lỗi, hết giờ, mất kết nối.", "UI"),
]


nfrs = [
    ("NFR-PERF-01", "API danh sách/chi tiết tài liệu phải phản hồi p95 <= 2 giây với dữ liệu demo, không tính tải tệp.", "PARTIAL"),
    ("NFR-PERF-02", "Upload phải trả kết quả tạo job trong <= 5 giây; indexing chạy nền và không khóa request.", "IMPLEMENTED"),
    ("NFR-PERF-03", "Chat mục tiêu phải phát token đầu tiên trong <= 5 giây và hoàn tất p95 <= 20 giây, phụ thuộc Gemini.", "NOT"),
    ("NFR-SEC-01", "Mọi API nghiệp vụ phải yêu cầu JWT hợp lệ; role phải được kiểm tra phía server.", "IMPLEMENTED"),
    ("NFR-SEC-02", "Secret, API key, password và raw embedding không được trả về client hoặc ghi log.", "PARTIAL"),
    ("NFR-SEC-03", "CORS chỉ cho phép origin cấu hình; không dùng wildcard cùng credentials trong production.", "PARTIAL"),
    ("NFR-SEC-04", "Upload phải kiểm tra MIME/magic bytes, kích thước, extension và chống path traversal/malware cơ bản.", "PARTIAL"),
    ("NFR-SEC-05", "JWT phải kiểm tra thuật toán, exp, issuer/audience khi triển khai production.", "PARTIAL"),
    ("NFR-REL-01", "Worker phải ghi trạng thái lỗi bền vững, không mất job khi tiến trình khởi động lại.", "IMPLEMENTED"),
    ("NFR-REL-02", "Các thao tác đa bước quan trọng phải dùng transaction hoặc cơ chế bù để tránh dữ liệu mồ côi.", "PARTIAL"),
    ("NFR-DATA-01", "Database phải duy trì FK, unique index và vector dimension thống nhất 3072.", "IMPLEMENTED"),
    ("NFR-DATA-02", "Backup/restore PostgreSQL và S3 phải được tài liệu hóa; mục tiêu RPO 24 giờ, RTO 4 giờ.", "PLANNED"),
    ("NFR-USE-01", "UI phải responsive tới 375px, hỗ trợ bàn phím, focus ring và tương phản WCAG AA.", "PARTIAL"),
    ("NFR-USE-02", "UI phải theo Mantine v7+ và DESIGN_SYSTEM.md; không dùng emoji làm icon hệ thống.", "PARTIAL"),
    ("NFR-MAINT-01", "Go backend phải giữ dependency hướng vào trong theo Clean Architecture; handler không truy cập DB trực tiếp.", "IMPLEMENTED"),
    ("NFR-MAINT-02", "Component/file nghiệp vụ nên dưới 200 dòng hoặc được tách hợp lý.", "PARTIAL"),
    ("NFR-TEST-01", "Core use case auth/document/chat phải có unit/integration tests và build phải qua CI.", "PARTIAL"),
    ("NFR-OBS-01", "Hệ thống phải có structured logging, correlation ID, health check và cảnh báo job lỗi.", "PARTIAL"),
    ("NFR-COMP-01", "Hỗ trợ phiên bản ổn định gần nhất của Chrome, Edge và Firefox.", "PLANNED"),
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("SWD392 Chatbot RAG | Software Requirements Specification")
    set_run(r, size=9, color=GRAY)
    footer = sec.footer.paragraphs[0]
    add_page_number(footer)

    # Editorial cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SOFTWARE REQUIREMENTS\nSPECIFICATION")
    set_run(r, size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SWD392 Chatbot RAG / StudyMate")
    set_run(r, size=18, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hệ thống quản lý học liệu, indexing và trợ lý học tập dựa trên RAG")
    set_run(r, size=12, italic=True, color=GRAY)
    doc.add_paragraph()
    add_table(doc, ["Thuộc tính", "Giá trị"], [
        ("Phiên bản", "1.0 - Baseline theo codebase"),
        ("Ngày lập", "13/06/2026"),
        ("Phạm vi khảo sát", "Toàn bộ repository: frontend, Go backend, Better Auth backend, database và docs"),
        ("Trạng thái", "Draft for review / Implementation-aware SRS"),
        ("Chuẩn tham chiếu", "IEEE/ISO/IEC 29148 theo cấu trúc thực dụng"),
    ], [2300, 7060], header_fill=LIGHT_GRAY, font_size=10)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Lưu ý quan trọng")
    set_run(r, size=11, bold=True, color=RED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tài liệu mô tả cả yêu cầu mục tiêu và hiện trạng. Mỗi yêu cầu đều có nhãn triển khai để tránh hiểu nhầm tính năng UI/mock là tính năng đã hoàn tất.")
    set_run(r, size=10, color=BLACK)
    doc.add_page_break()

    add_heading(doc, "Lịch sử phiên bản", 1)
    add_table(doc, ["Phiên bản", "Ngày", "Mô tả", "Tác giả"], [
        ("1.0", "13/06/2026", "Baseline SRS được tổng hợp từ codebase và tài liệu dự án.", "Codex"),
    ], [1200, 1500, 5160, 1500])

    add_heading(doc, "Quy ước trạng thái triển khai", 1)
    rows = [(label, meaning) for label, _, meaning in STATUS.values()]
    table = add_table(doc, ["Nhãn", "Ý nghĩa"], rows, [2200, 7160])
    for idx, key in enumerate(STATUS, start=1):
        shade(table.rows[idx].cells[0], STATUS[key][1])
        for run in table.rows[idx].cells[0].paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.bold = True

    add_heading(doc, "Mục lục nội dung", 1)
    for item in [
        "1. Giới thiệu", "2. Mô tả tổng quan", "3. Kiến trúc và môi trường",
        "4. Actor và phân quyền", "5. Yêu cầu chức năng", "6. Use case",
        "7. Yêu cầu dữ liệu", "8. Giao diện ngoài", "9. Yêu cầu phi chức năng",
        "10. Quy tắc nghiệp vụ", "11. Ngoài phạm vi", "12. Ma trận truy vết",
        "13. Kế hoạch nghiệm thu và khoảng trống", "Phụ lục A. Danh mục API hiện hữu",
        "Phụ lục B. Bằng chứng codebase",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_heading(doc, "1. Giới thiệu", 1)
    add_heading(doc, "1.1 Mục đích", 2)
    add_body(doc, "SRS này xác định các yêu cầu có thể kiểm thử cho SWD392 Chatbot RAG, đồng thời phân biệt rõ sản phẩm mục tiêu với phần đã triển khai trong repository tại ngày 13/06/2026. Tài liệu phục vụ thống nhất phạm vi, lập kế hoạch, phát triển, kiểm thử, demo và nghiệm thu.")
    add_heading(doc, "1.2 Phạm vi sản phẩm", 2)
    add_body(doc, "StudyMate là web app học thuật gồm quản lý người dùng và học liệu, upload/indexing tài liệu, tìm kiếm/đọc nội dung, kiểm duyệt, và trợ lý hỏi đáp RAG có trích dẫn. Codebase còn chứa định hướng quiz/practice và dashboard học tập; các phần này được giữ trong SRS nhưng gắn trạng thái riêng.")
    add_heading(doc, "1.3 Đối tượng đọc", 2)
    for text in ["Product owner và giảng viên hướng dẫn.", "Nhóm frontend, auth backend, Go/RAG backend và database.", "QA/tester và người thực hiện nghiệm thu.", "Người bảo trì, vận hành và đánh giá bảo mật."]:
        add_bullet(doc, text)
    add_heading(doc, "1.4 Thuật ngữ", 2)
    add_table(doc, ["Thuật ngữ", "Định nghĩa"], [
        ("RAG", "Retrieval-Augmented Generation: truy xuất ngữ cảnh trước khi sinh câu trả lời."),
        ("Chunk", "Đoạn văn bản nhỏ được tách từ tài liệu và gắn embedding."),
        ("Embedding", "Vector số biểu diễn ngữ nghĩa; hiện trạng dùng 3072 chiều."),
        ("Ground truth", "Câu trả lời chuẩn do con người chuẩn bị để đối chiếu."),
        ("Visibility", "Mức truy cập tài liệu: public, school_wide hoặc private."),
        ("RBAC", "Phân quyền dựa trên vai trò Admin/Lecturer/Student."),
    ], [1800, 7560])
    add_heading(doc, "1.5 Nguồn khảo sát", 2)
    for text in [
        "AGENTS.md và docs/system/DESIGN_SYSTEM.md.",
        "docs/Architecture.md, docs/ERD.txt, docs/Requirement.txt, docs/api_reference.md.",
        "Go router, handlers, application service, domain, repositories, worker, parsers, embedding và schema SQL.",
        "Better Auth/Hono service, JWT integration và admin user endpoints.",
        "Next.js routes, API clients, hooks, components và mock/local stores.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2. Mô tả tổng quan", 1)
    add_heading(doc, "2.1 Bối cảnh sản phẩm", 2)
    add_body(doc, "Hệ thống có kiến trúc tách dịch vụ: Next.js frontend trên cổng 3000; Hono/Better Auth trên cổng 5000; Go/Gin backend nghiệp vụ trên cổng 8080; PostgreSQL + pgvector; AWS S3 hoặc mock local cho tệp; Google Gemini cho embedding và phân chương. Frontend không được kết nối DB trực tiếp.")
    add_heading(doc, "2.2 Năng lực chính", 2)
    for text in [
        "Xác thực, JWT và RBAC ba vai trò.",
        "Quản lý người dùng, học kỳ, môn học, metadata và phân công Lecturer.",
        "Upload, lưu trữ, parse, chunk, embedding, pgvector và phân chương AI.",
        "Danh sách, tìm kiếm, đọc, sửa, xóa, duyệt và báo cáo tài liệu.",
        "Chat RAG có session/citation là năng lực mục tiêu nhưng chưa có backend thực.",
        "Quiz/practice và system settings hiện chủ yếu là UI prototype.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "2.3 Giả định và phụ thuộc", 2)
    for text in [
        "PostgreSQL đã bật extension pgvector và schema phù hợp vector(3072).",
        "JWT_SECRET/BETTER_AUTH_SECRET nhất quán giữa Hono và Go.",
        "Gemini API key hợp lệ và có quota.",
        "S3 bucket/credentials hợp lệ hoặc môi trường dev cho phép mock storage.",
        "Môn học, học kỳ, role và metadata cơ bản đã được seed.",
        "Không OCR tài liệu scan và không hỗ trợ realtime WebSocket trong baseline.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "2.4 Mâu thuẫn tài liệu đã được chuẩn hóa", 2)
    add_table(doc, ["Chủ đề", "Tài liệu cũ", "Baseline SRS"], [
        ("Embedding", "768 chiều", "3072 chiều theo code và SQL hiện hành."),
        ("Kiến trúc", "Go + Next monolith", "Ba dịch vụ tách rời: Next, Hono Auth, Go RAG."),
        ("Vai trò", "User đơn giản, không admin", "Admin/Lecturer/Student đã có trong code."),
        ("File type", "PDF/DOCX/PPTX/TXT/MD", "Upload handler hiện nhận PDF/DOC/DOCX/PPT/PPTX; TXT/MD chưa cam kết."),
        ("Chat", "Được ghi là hoàn tất", "Chưa có Go chat routes/retrieval; frontend dùng mock/local store."),
        ("Storage", "Local disk", "Code hiện ưu tiên AWS S3, có nhánh mock local."),
    ], [1700, 3400, 4260], font_size=8.5)

    add_heading(doc, "3. Kiến trúc và môi trường", 1)
    add_heading(doc, "3.1 Thành phần", 2)
    add_table(doc, ["Thành phần", "Công nghệ", "Trách nhiệm", "Trạng thái"], [
        ("Frontend", "Next.js 15, React 19, Mantine 7", "UI, route guard, API clients", "Đã có; một số màn hình mock"),
        ("Auth backend", "Hono, Better Auth, JOSE", "Đăng ký/đăng nhập/session/JWT/user admin", "Đã có"),
        ("RAG backend", "Go 1.25, Gin", "Documents, admin, worker, embedding", "Đã có; thiếu chat query"),
        ("Database", "PostgreSQL + pgvector", "Users, documents, chunks, metadata, jobs", "Đã có schema"),
        ("Object storage", "AWS S3", "Tệp gốc", "Đã tích hợp; có mock dev"),
        ("AI", "Gemini Embedding 2 / 2.5 Flash", "Embedding và phân chương", "Đã có"),
    ], [1450, 1900, 3560, 2450], font_size=8.3)
    add_heading(doc, "3.2 Ràng buộc kiến trúc", 2)
    for text in [
        "Go tuân thủ Clean Architecture; dependency hướng vào domain/application.",
        "Handler không truy cập DB trực tiếp; mọi nghiệp vụ qua service/use case/repository.",
        "Frontend chỉ gọi Hono/Go API, không truy cập PostgreSQL.",
        "Cấu hình qua environment variables; không hardcode secret trong production.",
        "Request/response HTTP chuẩn; không WebSocket trong baseline.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "4. Actor và phân quyền", 1)
    add_table(doc, ["Actor", "Mô tả", "Quyền chính"], [
        ("Visitor", "Chưa xác thực", "Đăng ký, đăng nhập, yêu cầu đặt lại mật khẩu."),
        ("Student", "Người học", "Xem/đọc/report tài liệu; chat/quiz mục tiêu."),
        ("Lecturer", "Giảng viên", "Quyền Student + upload và quản lý tài liệu cá nhân; tạo quiz mục tiêu."),
        ("Admin", "Quản trị", "Users, curriculum, assignment, moderation, system settings."),
        ("Worker", "Tiến trình nền", "Parse, chunk, embed, lưu vector, phân chương, cập nhật job."),
        ("Gemini", "Dịch vụ AI ngoài", "Embedding, phân chương và chat generation mục tiêu."),
    ], [1400, 2600, 5360])

    add_heading(doc, "5. Yêu cầu chức năng", 1)
    add_body(doc, "Từ khóa “phải” biểu thị yêu cầu bắt buộc của sản phẩm mục tiêu. Nhãn trạng thái cho biết mức độ hiện thực hóa tại thời điểm khảo sát, không làm giảm tính bắt buộc của yêu cầu.")
    groups = [
        ("5.1 Xác thực và phân quyền", ("FR-AUTH", "FR-RBAC")),
        ("5.2 Người dùng và học vụ", ("FR-USER", "FR-CUR")),
        ("5.3 Tài liệu và indexing", ("FR-DOC", "FR-IDX")),
        ("5.4 Khám phá và kiểm duyệt", ("FR-DISC", "FR-MOD")),
        ("5.5 Chat và RAG", ("FR-CHAT", "FR-RAG")),
        ("5.6 Quiz và practice", ("FR-QUIZ",)),
        ("5.7 Settings, audit và đánh giá", ("FR-SET", "FR-AUD", "FR-EVAL")),
    ]
    for heading, prefixes in groups:
        add_heading(doc, heading, 2)
        for req in requirements:
            if req["id"].startswith(prefixes):
                add_requirement(doc, req)

    add_heading(doc, "6. Use case", 1)
    for uc_id, name, actor, pre, main, alt, status in use_cases:
        add_heading(doc, f"{uc_id} - {name}", 2)
        add_status_badge(doc, status, STATUS[status][2])
        add_body(doc, f"Actor: {actor}", bold_prefix="Actor:")
        add_body(doc, f"Tiền điều kiện: {pre}", bold_prefix="Tiền điều kiện:")
        add_body(doc, f"Luồng chính: {main}", bold_prefix="Luồng chính:")
        add_body(doc, f"Ngoại lệ: {alt}", bold_prefix="Ngoại lệ:")
        add_body(doc, "Hậu điều kiện: Dữ liệu và trạng thái liên quan được lưu nhất quán; người dùng nhận phản hồi rõ ràng.")

    add_heading(doc, "7. Yêu cầu dữ liệu", 1)
    add_heading(doc, "7.1 Mô hình dữ liệu baseline", 2)
    add_table(doc, ["Nhóm", "Bảng/thực thể", "Mục đích", "Trạng thái"], [
        ("Auth", "users, roles, account, session, verification", "Tài khoản, credential và phiên Better Auth", "Hiện hữu"),
        ("Học vụ", "academic_terms, subjects, user_subjects", "Học kỳ, môn học, phân công Lecturer", "Hiện hữu"),
        ("Tài liệu", "documents, document_files, document_chapters", "Metadata, file và cấu trúc chương", "Hiện hữu"),
        ("Vector", "document_chunks", "Chunk nội dung và embedding vector(3072)", "Hiện hữu"),
        ("Vận hành", "upload_jobs, audit_logs, document_reports", "Job, audit, moderation", "Hiện hữu"),
        ("Cá nhân", "user_bookmarks", "Bookmark tài liệu", "Schema-only"),
        ("Chat", "chat_sessions, messages, message_citations", "Lịch sử và citation", "Chưa có trong schema hiện hành"),
        ("Quiz", "quizzes, questions, attempts, answers", "Quiz/practice", "Chưa có"),
    ], [1150, 2600, 3710, 1900], font_size=8.3)
    add_heading(doc, "7.2 Quy tắc dữ liệu", 2)
    for text in [
        "UUID là khóa chính mặc định cho dữ liệu nghiệp vụ; role dùng smallint 1/2/3.",
        "Email, username, role name, subject code và các code danh mục phải duy nhất theo schema.",
        "Embedding phải thống nhất 3072 chiều giữa model, code và pgvector.",
        "Xóa document phải xử lý file, chunks, chapters, reports, bookmarks và jobs theo FK/chính sách.",
        "Password, token, API key và embedding thô không xuất hiện trong response nghiệp vụ.",
        "Dữ liệu chat và quiz khi triển khai phải có owner, timestamps, retention và cascade rõ ràng.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "7.3 Retention và riêng tư", 2)
    add_body(doc, "Baseline chưa định nghĩa đầy đủ retention. Yêu cầu mục tiêu: session auth hết hạn theo cấu hình; audit log tối thiểu 180 ngày; upload job tối thiểu 90 ngày; chat/quiz tồn tại đến khi người dùng hoặc chính sách xóa; tệp và vector phải được xóa đồng bộ khi document bị xóa.")

    add_heading(doc, "8. Giao diện ngoài", 1)
    add_heading(doc, "8.1 Giao diện người dùng", 2)
    for text in [
        "Mantine v7+, Plus Jakarta Sans cho UI, Newsreader cho nội dung học thuật và Tabler Icons.",
        "Ba shell theo role; route được bảo vệ và chuyển hướng đúng home.",
        "Bảng dữ liệu responsive; form có validation, loading và thông báo lỗi.",
        "Chat citation phải có khả năng mở đoạn nguồn; UI hiện mới mô phỏng.",
        "Hỗ trợ màn hình 375px trở lên, bàn phím và focus rõ.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "8.2 API và giao thức", 2)
    add_body(doc, "API dùng JSON qua HTTP/HTTPS; upload dùng multipart/form-data; bearer JWT trong Authorization. Auth cookie có thể dùng cho Better Auth session, nhưng Go API nhận JWT. Production phải dùng HTTPS và origin CORS cấu hình.")
    add_heading(doc, "8.3 Dịch vụ bên ngoài", 2)
    add_table(doc, ["Dịch vụ", "Mục đích", "Failure handling"], [
        ("Gemini Embedding", "Sinh vector 3072", "Timeout, retry/backoff, rotate API keys, fail job."),
        ("Gemini Generate", "Phân chương; chat mục tiêu", "Fallback chương; chat trả lỗi có thể thử lại."),
        ("AWS S3", "Lưu tệp gốc", "Bù/xóa document khi upload lỗi; retry theo SDK."),
        ("PostgreSQL", "Dữ liệu và vector", "Health check, transaction, timeout query."),
    ], [1900, 3000, 4460])

    add_heading(doc, "9. Yêu cầu phi chức năng", 1)
    add_table(doc, ["ID", "Yêu cầu", "Hiện trạng"], [(a, b, STATUS[c][0]) for a, b, c in nfrs], [1500, 6260, 1600], font_size=8.5)

    add_heading(doc, "10. Quy tắc nghiệp vụ", 1)
    rules = [
        ("BR-01", "role_id 1=Admin, 2=Lecturer, 3=Student."),
        ("BR-02", "Chỉ Lecturer được upload và quản lý tài liệu cá nhân."),
        ("BR-03", "Student/Lecturer chỉ xem tài liệu theo visibility, status và ownership."),
        ("BR-04", "Admin có quyền duyệt/reject/delete tài liệu và xử lý report."),
        ("BR-05", "Tài liệu chỉ được dùng cho RAG khi indexing hoàn tất và được phép truy cập."),
        ("BR-06", "Mỗi chunk thuộc đúng một document và có chunk_order duy nhất trong document."),
        ("BR-07", "Mỗi user-subject và bookmark là một cặp duy nhất."),
        ("BR-08", "Chat không đủ context phải đánh dấu out-of-scope và không bịa câu trả lời."),
        ("BR-09", "Citation phải trỏ tới chunk đã dùng trong retrieval."),
        ("BR-10", "Benchmark/RAGAS tự động không thuộc baseline; bộ 50 câu ground truth vẫn là deliverable."),
    ]
    add_table(doc, ["ID", "Quy tắc"], rules, [1300, 8060])

    add_heading(doc, "11. Ngoài phạm vi", 1)
    for text in [
        "OCR ảnh/scan, nhận dạng chữ viết tay và xử lý media nặng.",
        "Realtime collaboration, WebSocket chat hoặc voice conversation hoàn chỉnh.",
        "Multi-tenancy giữa nhiều trường/tổ chức và billing/subscription.",
        "RAGAS/experiment management hoặc nghiên cứu benchmark tự động quy mô lớn.",
        "Mobile app native, offline mode và đồng bộ đa thiết bị nâng cao.",
        "SSO enterprise, social login và 2FA trong baseline (UI settings có thể đề cập nhưng chưa cam kết).",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "12. Ma trận truy vết", 1)
    trace_rows = []
    mapping = {
        "AUTH": "UC-01", "RBAC": "UC-01", "USER": "UC-06", "CUR": "UC-06",
        "DOC": "UC-02/03/04", "IDX": "UC-02", "DISC": "UC-03", "MOD": "UC-05",
        "CHAT": "UC-07", "RAG": "UC-07", "QUIZ": "UC-08", "SET": "UC-06",
        "AUD": "UC-05/06", "EVAL": "UC-07",
    }
    for req in requirements:
        family = req["id"].split("-")[1]
        tests = "API/Integration" if req["status"] in ("IMPLEMENTED", "PARTIAL") else "Acceptance/Future"
        trace_rows.append((req["id"], mapping.get(family, "-"), STATUS[req["status"]][0], tests))
    add_table(doc, ["Requirement", "Use case", "Trạng thái", "Loại test"], trace_rows, [1900, 1800, 3100, 2560], font_size=8.2)

    add_heading(doc, "13. Kế hoạch nghiệm thu và khoảng trống", 1)
    add_heading(doc, "13.1 Điều kiện nghiệm thu tối thiểu", 2)
    for text in [
        "Ba service build và khởi động bằng cấu hình mẫu; health endpoints phản hồi.",
        "Auth/RBAC được kiểm thử cho cả ba role và tài khoản bị khóa.",
        "Upload một PDF, DOCX và PPTX; job hoàn tất; chunks vector(3072) được lưu.",
        "Tìm kiếm/đọc/sửa/xóa/duyệt/report tài liệu hoạt động với quyền đúng.",
        "Chat RAG thật thay thế mock: session persistence, semantic retrieval, Gemini answer và citation.",
        "Bộ 50 câu ground truth tồn tại và có báo cáo chạy.",
        "Quiz/practice chỉ được tính vào nghiệm thu nếu backend/schema/API được bổ sung; UI mock không được coi là hoàn tất.",
        "Go tests, frontend lint/build và auth TypeScript check đều qua.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "13.2 Khoảng trống ưu tiên", 2)
    add_table(doc, ["Ưu tiên", "Khoảng trống", "Tác động"], [
        ("P0", "Không có backend Chat/RAG retrieval/generation/session/citation.", "Thiếu năng lực cốt lõi của sản phẩm."),
        ("P0", "Chưa có test set 50 câu + ground truth.", "Thiếu deliverable và cơ sở đánh giá."),
        ("P1", "Frontend chat/session/quiz/practice dùng mock/local state.", "Demo có thể gây hiểu nhầm là đã tích hợp."),
        ("P1", "Auth role có fallback suy từ email và token lưu localStorage.", "Rủi ro bảo mật/không nhất quán."),
        ("P1", "CORS wildcard ở Go và secret mặc định trong config.", "Không phù hợp production."),
        ("P1", "File type docs/code lệch; DOC/PPT cũ chưa rõ parser.", "Upload có thể nhận nhưng worker thất bại."),
        ("P2", "Bookmark và settings chỉ schema/UI.", "Tính năng phụ chưa bền vững."),
        ("P2", "Coverage test tập trung parser/chunker/embedding.", "Thiếu test service, repository, handler và auth."),
    ], [1000, 5060, 3300], font_size=8.5)

    add_heading(doc, "Phụ lục A. Danh mục API hiện hữu", 1)
    api_rows = [
        ("GET", "/api/health", "Public", "Health"),
        ("GET", "/api/documents/lookups", "All roles", "Metadata lookups"),
        ("GET", "/api/documents", "Lecturer, Student", "List/search documents"),
        ("GET", "/api/documents/:slug", "All roles", "Document details"),
        ("POST", "/api/documents/:slug/report", "Lecturer, Student", "Report document"),
        ("GET", "/api/documents/my", "Lecturer", "Owned documents"),
        ("GET", "/api/documents/dashboard", "Lecturer", "Dashboard/jobs"),
        ("POST", "/api/documents/upload", "Lecturer", "Upload"),
        ("POST", "/api/documents/:slug/edit", "Lecturer", "Edit metadata"),
        ("GET", "/api/documents/:slug/delete-view", "Lecturer", "Delete impact"),
        ("POST", "/api/documents/:slug/delete", "Lecturer", "Delete"),
        ("GET/POST", "/api/admin/users...", "Admin", "List/block/unblock"),
        ("GET/PUT", "/api/admin/.../subjects", "Admin", "Lecturer assignment"),
        ("GET/POST", "/api/admin/documents...", "Admin", "Review documents"),
        ("GET/POST", "/api/admin/reports...", "Admin", "Resolve reports"),
        ("POST/PUT/DELETE", "/api/admin/{metadata}", "Admin", "Metadata CRUD"),
        ("GET/POST", "/api/auth/*", "Public/session", "Better Auth"),
        ("POST/PUT/DELETE", "/api/admin/users...", "Admin via Hono", "Create/update/password/delete"),
    ]
    add_table(doc, ["Method", "Path", "Quyền", "Mục đích"], api_rows, [1300, 3600, 2100, 2360], font_size=8.2)
    add_body(doc, "Không tìm thấy endpoint Go cho /api/chat, semantic query, session/message persistence, citation hoặc quiz.")

    add_heading(doc, "Phụ lục B. Bằng chứng codebase", 1)
    evidence = [
        ("Go routes", "backend/go/internal/interface/router/router.go"),
        ("Auth/JWT", "backend/better-auth/auth.ts; backend/go/internal/interface/middleware/auth.go"),
        ("Upload/indexing", "document-handler.go; background_worker.go; embedding/gemini-embedding.go"),
        ("Database", "backend/database/swd391_dangerous_malware.sql"),
        ("Document service", "backend/go/internal/application/document_service.go"),
        ("Frontend API", "frontend/src/api/*.ts"),
        ("Mock chat/session", "frontend/src/components/common/chat/chat-view.tsx; frontend/src/lib/sessions-store.ts"),
        ("Mock quiz/practice", "frontend/src/hooks/student/use-quiz.ts; lecturer/use-create-quiz.ts; use-practice.ts"),
        ("Design rules", "docs/system/DESIGN_SYSTEM.md"),
        ("Legacy specs", "docs/Architecture.md; docs/ERD.txt; docs/Requirement.txt; docs/planing/"),
    ]
    add_table(doc, ["Khu vực", "Đường dẫn"], evidence, [2400, 6960])
    add_body(doc, "Kết luận baseline: quản lý học liệu và pipeline indexing đã có nền tảng đáng kể; phần hỏi đáp RAG end-to-end vẫn là khoảng trống cốt lõi. Mọi demo/review phải dùng nhãn trạng thái trong SRS này để tránh đánh đồng UI prototype với tính năng production.")

    # Keep headings from orphaning and set table font consistently.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.1

    doc.core_properties.title = "Software Requirements Specification - SWD392 Chatbot RAG"
    doc.core_properties.subject = "Implementation-aware SRS"
    doc.core_properties.author = "SWD392 Team"
    doc.core_properties.keywords = "SRS, RAG, StudyMate, SWD392, requirements"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
